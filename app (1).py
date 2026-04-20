"""
EmpiricX — Research Intelligence Engine
Single-file Streamlit app. Deploy directly to Streamlit Community Cloud.

Requirements (requirements.txt):
    streamlit>=1.32.0
    openai>=1.30.0
    PyMuPDF>=1.24.0
    python-docx>=1.1.0
    openpyxl>=3.1.2
    pandas>=2.0.0
"""

# ═══════════════════════════════════════════════════════════════════════
# IMPORTS
# ═══════════════════════════════════════════════════════════════════════
import io
import os
import re
import csv
import json
import time
import datetime
import pandas as pd
import streamlit as st
from collections import Counter
from typing import Any, Optional


# ═══════════════════════════════════════════════════════════════════════
# PAGE CONFIG  (must be first Streamlit call)
# ═══════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="EmpiricX — Research Intelligence",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ═══════════════════════════════════════════════════════════════════════
# GLOBAL CSS
# ═══════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Lora:ital,wght@0,400;0,600;0,700;1,400;1,600&family=Outfit:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --bg:#07070d;--bg-2:#0c0c18;--bg-3:#111120;
  --panel:#151526;--panel-2:#1b1b2f;--panel-3:#20203a;
  --gold:#d4a843;--gold-lt:#f0c866;--gold-dk:#9a7a2a;
  --gold-glow:rgba(212,168,67,.18);--gold-soft:rgba(212,168,67,.08);--gold-mid:rgba(212,168,67,.20);
  --blue:#5585ff;--blue-soft:rgba(85,133,255,.1);
  --green:#3ecf8e;--green-soft:rgba(62,207,142,.1);
  --red:#f05959;--red-soft:rgba(240,89,89,.1);
  --violet:#9b7eff;--violet-soft:rgba(155,126,255,.1);
  --t1:#f0f0f8;--t2:#9898b8;--t3:#52526a;--t4:#36364a;
  --b1:rgba(255,255,255,.055);--b2:rgba(255,255,255,.10);--b3:rgba(255,255,255,.16);
  --serif:'Lora',Georgia,serif;
  --sans:'Outfit',system-ui,sans-serif;
  --mono:'JetBrains Mono',monospace;
  --r:10px;--r-lg:16px;--r-xl:22px;
  --t-xs:clamp(.68rem,1.6vw,.78rem);--t-sm:clamp(.84rem,2vw,.95rem);
  --t-base:clamp(.96rem,2.4vw,1.05rem);--t-md:clamp(1.1rem,2.8vw,1.2rem);
  --t-lg:clamp(1.35rem,3.5vw,1.65rem);--t-xl:clamp(1.8rem,5vw,2.5rem);
  --t-2xl:clamp(2.2rem,6.5vw,3.1rem);
}

*,*::before,*::after{box-sizing:border-box;}
html{font-size:16px;-webkit-text-size-adjust:100%;}
html,body,[data-testid="stAppViewContainer"],[data-testid="stApp"]{
  background:var(--bg)!important;color:var(--t1)!important;
  font-family:var(--sans)!important;font-size:var(--t-base)!important;line-height:1.65!important;}
footer,[data-testid="stDecoration"]{display:none!important;}
[data-testid="stHeader"]{background:rgba(7,7,13,.9)!important;backdrop-filter:blur(12px)!important;border-bottom:1px solid var(--b1)!important;}
[data-testid="stHeader"] button{color:var(--t3)!important;background:transparent!important;}
[data-testid="stHeader"] button:hover{color:var(--gold)!important;background:var(--gold-soft)!important;}
[data-testid="stMainBlockContainer"]{padding:1.75rem clamp(1rem,4.5vw,3rem) 4rem!important;max-width:100%!important;}
@media(min-width:1200px){[data-testid="stMainBlockContainer"]{padding:2.25rem clamp(2rem,5.5vw,4.5rem) 5rem!important;}}
[data-testid="stSidebar"]{background:var(--bg-2)!important;border-right:1px solid var(--b1)!important;}
[data-testid="stSidebar"]>div:first-child{padding:1.5rem 1.25rem 2rem!important;}

/* ── Sidebar brand ── */
.sb-brand{display:flex;align-items:center;gap:12px;padding:.2rem 0 .6rem;}
.sb-logomark{width:36px;height:36px;flex-shrink:0;display:block;}
.sb-name{font-family:var(--serif);font-size:1.35rem;font-weight:700;color:var(--t1);letter-spacing:-.02em;}
.sb-name em{color:var(--gold);font-style:italic;}
.sb-tag{font-family:var(--mono);font-size:.58rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:1.1rem;padding-left:48px;}
.sb-divider{height:1px;background:var(--b1);margin:.85rem 0;}
.sb-label{font-family:var(--mono)!important;font-size:.6rem!important;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3)!important;margin-bottom:10px!important;display:block;}
[data-testid="stSidebar"] [data-testid="stFileUploader"]{background:var(--bg-3)!important;border:1.5px dashed var(--b2)!important;border-radius:var(--r-lg)!important;transition:border-color .2s!important;}
[data-testid="stSidebar"] [data-testid="stFileUploader"]:hover{border-color:var(--gold)!important;}
[data-testid="stSidebar"] [data-testid="stFileUploader"] section{padding:1.1rem 1rem!important;}
[data-testid="stSidebar"] [data-testid="stFileUploader"] label,[data-testid="stSidebar"] [data-testid="stFileUploader"] p,[data-testid="stSidebar"] [data-testid="stFileUploader"] span{color:var(--t2)!important;font-size:.82rem!important;font-family:var(--sans)!important;}
.sb-paper{display:flex;align-items:flex-start;gap:9px;background:var(--bg-3);border:1px solid var(--b1);border-radius:var(--r);padding:10px 12px;margin-bottom:6px;transition:border-color .15s;}
.sb-paper:hover{border-color:var(--b2);}
.sb-paper-icon{width:28px;height:28px;border-radius:7px;background:var(--panel-2);border:1px solid var(--b1);display:flex;align-items:center;justify-content:center;font-size:.9rem;flex-shrink:0;margin-top:1px;}
.sb-paper-name{font-size:.82rem;font-weight:500;color:var(--t1);line-height:1.3;word-break:break-word;}
.sb-paper-meta{font-family:var(--mono);font-size:.64rem;color:var(--t3);margin-top:3px;}
.sb-badge{margin-left:auto;flex-shrink:0;font-family:var(--mono);font-size:.6rem;letter-spacing:.5px;padding:3px 9px;border-radius:20px;white-space:nowrap;}
.sb-badge-ok{background:var(--green-soft);color:var(--green);border:1px solid rgba(62,207,142,.25);}
.sb-badge-new{background:var(--blue-soft);color:var(--blue);border:1px solid rgba(85,133,255,.25);}
.sb-stats{display:flex;gap:8px;margin:.6rem 0;}
.sb-stat{flex:1;background:var(--bg-3);border:1px solid var(--b1);border-radius:var(--r);padding:12px 8px;text-align:center;}
.sb-stat-n{font-family:var(--serif);font-size:1.7rem;color:var(--gold);line-height:1;}
.sb-stat-l{font-family:var(--mono);font-size:.58rem;letter-spacing:1.5px;text-transform:uppercase;color:var(--t3);margin-top:4px;}
[data-testid="stRadio"] label{font-family:var(--sans)!important;font-size:.92rem!important;color:var(--t2)!important;padding:9px 14px!important;border-radius:var(--r)!important;border:1px solid transparent!important;transition:all .15s!important;cursor:pointer!important;line-height:1.45!important;}
[data-testid="stRadio"] label:hover{color:var(--t1)!important;background:var(--panel)!important;border-color:var(--b1)!important;}

/* ── Buttons ── */
.stButton>button{background:linear-gradient(135deg,#d4a843 0%,#a87e2a 100%)!important;color:#0a0700!important;border:none!important;border-radius:var(--r)!important;font-family:var(--sans)!important;font-weight:600!important;font-size:.92rem!important;letter-spacing:.1px!important;padding:.65rem 1.6rem!important;min-height:2.65rem!important;transition:all .2s cubic-bezier(.16,1,.3,1)!important;box-shadow:0 2px 12px var(--gold-glow)!important;}
.stButton>button:hover{transform:translateY(-2px)!important;box-shadow:0 8px 28px rgba(212,168,67,.32)!important;background:linear-gradient(135deg,#e8bc55 0%,#bc9035 100%)!important;}
.stButton>button:active{transform:translateY(0) scale(.98)!important;}
[data-testid="stBaseButton-secondary"]>button{background:transparent!important;border:1px solid var(--b2)!important;color:var(--t2)!important;box-shadow:none!important;}
[data-testid="stBaseButton-secondary"]>button:hover{border-color:var(--gold)!important;color:var(--gold)!important;background:var(--gold-soft)!important;transform:none!important;box-shadow:none!important;}
[data-testid="stDownloadButton"]>button{min-height:2.65rem!important;font-size:.92rem!important;}

/* ── Inputs ── */
.stTextInput input,[data-baseweb="input"] input{background:var(--bg-3)!important;border:1.5px solid var(--b2)!important;border-radius:var(--r)!important;color:var(--t1)!important;font-family:var(--sans)!important;font-size:1rem!important;padding:.7rem 1.1rem!important;min-height:2.9rem!important;transition:border-color .2s,box-shadow .2s!important;}
.stTextInput input:focus,[data-baseweb="input"] input:focus{border-color:var(--gold)!important;box-shadow:0 0 0 3px var(--gold-glow)!important;outline:none!important;}
.stTextInput label{font-family:var(--mono)!important;font-size:.66rem!important;letter-spacing:2.5px!important;text-transform:uppercase!important;color:var(--t3)!important;margin-bottom:8px!important;}

/* ── Page headers ── */
.ph-wrap{padding:1.5rem 0 1.6rem;border-bottom:1px solid var(--b1);margin-bottom:1.75rem;position:relative;}
.ph-wrap::after{content:'';position:absolute;bottom:-1px;left:0;width:64px;height:2px;background:linear-gradient(90deg,var(--gold),transparent);}
.ph-eye{display:inline-flex;align-items:center;gap:7px;font-family:var(--mono);font-size:.66rem;letter-spacing:3px;text-transform:uppercase;color:var(--gold);margin-bottom:11px;background:var(--gold-soft);border:1px solid var(--gold-mid);border-radius:20px;padding:4px 12px 4px 10px;}
.ph-eye::before{content:'\25C6';font-size:.45rem;opacity:.7;}
.ph-title{font-family:var(--serif);font-size:var(--t-xl);color:var(--t1);line-height:1.1;margin:0;font-weight:700;}
.ph-title span{color:var(--gold);font-style:italic;}
.ph-sub{font-family:var(--sans);font-size:.97rem;color:var(--t2);margin-top:12px;max-width:580px;line-height:1.7;}

/* ── Gate / Landing ── */
.gate-bg{position:fixed;inset:0;z-index:0;background:radial-gradient(ellipse 80% 60% at 15% 110%,rgba(212,168,67,.06) 0%,transparent 60%),radial-gradient(ellipse 60% 50% at 88% -5%,rgba(85,133,255,.05) 0%,transparent 55%),var(--bg);pointer-events:none;}
.gate-bg::before{content:'';position:absolute;inset:0;background-image:linear-gradient(rgba(255,255,255,.018) 1px,transparent 1px),linear-gradient(90deg,rgba(255,255,255,.018) 1px,transparent 1px);background-size:56px 56px;mask-image:radial-gradient(ellipse 75% 75% at 50% 50%,black,transparent);animation:gridDrift 25s linear infinite;}
@keyframes gridDrift{0%{transform:translate(0,0);}100%{transform:translate(56px,56px);}}
.gate-card{position:relative;z-index:1;background:rgba(18,18,30,.92);border:1px solid var(--b2);border-radius:var(--r-xl);padding:clamp(2rem,6vw,3.25rem) clamp(1.75rem,6vw,3.75rem);width:100%;backdrop-filter:blur(24px);box-shadow:0 40px 80px rgba(0,0,0,.55),0 0 0 1px rgba(212,168,67,.05);animation:cardRise .55s cubic-bezier(.16,1,.3,1) both;}
@keyframes cardRise{from{opacity:0;transform:translateY(24px) scale(.97);}to{opacity:1;transform:translateY(0) scale(1);}}
.gate-logo{text-align:center;margin-bottom:.35rem;font-family:var(--serif);font-size:clamp(2rem,7.5vw,2.8rem);font-weight:700;color:var(--t1);}
.gate-logo span{color:var(--gold);font-style:italic;}
.gate-sub{text-align:center;margin-bottom:clamp(1rem,3vw,1.5rem);font-family:var(--mono);font-size:.64rem;letter-spacing:3px;text-transform:uppercase;color:var(--t3);}
.gate-pill{display:flex;align-items:center;justify-content:center;gap:7px;background:var(--gold-soft);border:1px solid var(--gold-mid);border-radius:20px;padding:6px 16px;font-family:var(--mono);font-size:.64rem;color:var(--gold);letter-spacing:1px;margin:0 auto clamp(1.5rem,4vw,2rem);width:fit-content;}
.gate-features{display:flex;flex-direction:column;gap:10px;margin-bottom:clamp(1.5rem,4vw,2.25rem);}
.gate-feat{display:flex;align-items:flex-start;gap:11px;font-family:var(--sans);font-size:.9rem;color:var(--t2);line-height:1.5;}
.gate-feat-dot{width:6px;height:6px;border-radius:50%;background:var(--gold);flex-shrink:0;margin-top:8px;}
.gate-error{background:var(--red-soft);border:1px solid rgba(240,89,89,.28);border-radius:var(--r);padding:11px 16px;font-family:var(--sans);font-size:.88rem;color:var(--red);margin-bottom:1rem;text-align:center;}
.gate-footer{text-align:center;margin-top:clamp(1rem,3vw,1.75rem);font-family:var(--sans);font-size:.74rem;color:var(--t3);}
.gate-footer strong{color:var(--gold);font-weight:500;}

/* ── Pricing cards ── */
.pricing-section{margin:clamp(2rem,5vw,3rem) 0 clamp(1.5rem,4vw,2.5rem);}
.pricing-section-title{font-family:var(--serif);font-size:clamp(1.1rem,3vw,1.4rem);font-weight:700;color:var(--t1);text-align:center;margin-bottom:.4rem;}
.pricing-section-sub{font-family:var(--sans);font-size:.85rem;color:var(--t3);text-align:center;margin-bottom:clamp(1.2rem,3.5vw,2rem);}
.pricing-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;}
@media(max-width:700px){.pricing-grid{grid-template-columns:1fr;gap:12px;}}
.pricing-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.5rem 1.4rem;position:relative;overflow:hidden;transition:transform .2s,border-color .2s,box-shadow .2s;display:flex;flex-direction:column;}
.pricing-card:hover{transform:translateY(-4px);border-color:var(--b2);box-shadow:0 16px 40px rgba(0,0,0,.35);}
.pricing-card-popular{border-color:var(--gold)!important;box-shadow:0 0 0 1px var(--gold-mid),0 16px 48px rgba(212,168,67,.12);}
.pricing-badge{position:absolute;top:0;right:0;background:linear-gradient(135deg,var(--gold),var(--gold-dk));color:#07070d;font-family:var(--mono);font-size:.55rem;letter-spacing:2px;text-transform:uppercase;font-weight:700;padding:5px 14px;border-radius:0 var(--r-lg) 0 var(--r-lg);}
.pricing-tier{font-family:var(--mono);font-size:.6rem;letter-spacing:3px;text-transform:uppercase;color:var(--t3);margin-bottom:10px;}
.pricing-price{font-family:var(--serif);font-size:clamp(1.9rem,5vw,2.6rem);color:var(--t1);line-height:1;margin-bottom:4px;}
.pricing-price span{font-size:.85rem;font-family:var(--sans);color:var(--t3);font-weight:400;}
.pricing-desc{font-family:var(--sans);font-size:.8rem;color:var(--t3);margin-bottom:1.25rem;line-height:1.5;}
.pricing-features{list-style:none;padding:0;margin:0 0 1.5rem;flex:1;}
.pricing-features li{display:flex;align-items:flex-start;gap:9px;font-family:var(--sans);font-size:.82rem;color:var(--t2);padding:7px 0;border-bottom:1px solid var(--b1);line-height:1.45;}
.pricing-features li:last-child{border-bottom:none;}
.pricing-features li::before{content:'✦';color:var(--gold);font-size:.55rem;flex-shrink:0;margin-top:5px;}
.pricing-cta{display:block;width:100%;padding:.7rem 1rem;border-radius:var(--r);font-family:var(--sans);font-size:.9rem;font-weight:600;text-align:center;cursor:pointer;transition:all .2s;text-decoration:none;border:none;}
.pricing-cta-primary{background:linear-gradient(135deg,#d4a843 0%,#a87e2a 100%);color:#0a0700;box-shadow:0 4px 18px var(--gold-glow);}
.pricing-cta-primary:hover{background:linear-gradient(135deg,#e8bc55 0%,#bc9035 100%);box-shadow:0 8px 28px rgba(212,168,67,.32);transform:translateY(-2px);}
.pricing-cta-outline{background:transparent;color:var(--t2);border:1px solid var(--b2)!important;}
.pricing-cta-outline:hover{border-color:var(--gold)!important;color:var(--gold);background:var(--gold-soft);}

/* ── Gate footer links ── */
.gate-links{display:flex;justify-content:center;align-items:center;gap:20px;margin-top:1.1rem;flex-wrap:wrap;}
.gate-link{font-family:var(--sans);font-size:.78rem;color:var(--t3);text-decoration:none;display:inline-flex;align-items:center;gap:5px;padding:5px 10px;border-radius:var(--r);border:1px solid transparent;transition:all .15s;cursor:pointer;}
.gate-link:hover{color:var(--gold);border-color:var(--gold-mid);background:var(--gold-soft);}
.gate-link-dot{width:3px;height:3px;border-radius:50%;background:var(--b2);}

/* ── Metric chips ── */
.m-row{display:flex;gap:14px;flex-wrap:wrap;margin-bottom:2rem;}
.m-chip{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.2rem 1.5rem;flex:1;min-width:120px;position:relative;overflow:hidden;transition:border-color .2s,box-shadow .2s;}
.m-chip::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,var(--gold),transparent);opacity:.5;}
.m-chip:hover{border-color:var(--b2);box-shadow:0 4px 20px rgba(0,0,0,.28);}
.m-val{font-family:var(--serif);font-size:clamp(1.6rem,4.5vw,2.1rem);color:var(--gold);line-height:1;margin-bottom:5px;}
.m-lbl{font-family:var(--mono);font-size:.62rem;letter-spacing:1.5px;text-transform:uppercase;color:var(--t3);}

/* ── Cards ── */
.x-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:clamp(1.1rem,3vw,1.5rem);margin-bottom:1rem;transition:border-color .2s,box-shadow .2s;}
.x-card:hover{border-color:var(--b2);box-shadow:0 4px 24px rgba(0,0,0,.25);}
.x-card-gold{border-left:3px solid var(--gold);}
.detail-block{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r);padding:1rem 1.2rem;margin-bottom:.75rem;transition:border-color .15s;}
.detail-block:hover{border-color:var(--b2);}
.detail-lbl{font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:7px;}
.detail-val{font-family:var(--sans);font-size:.93rem;color:var(--t2);line-height:1.65;}

/* ── Synthesis prose ── */
.syn-overview{background:linear-gradient(135deg,var(--panel) 0%,var(--panel-2) 100%);border:1px solid var(--b2);border-left:4px solid var(--gold);border-radius:var(--r-lg);padding:clamp(1.4rem,4vw,2rem) clamp(1.5rem,4vw,2.2rem);margin-bottom:1.75rem;position:relative;}
.syn-overview-label{font-family:var(--mono);font-size:.6rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--gold);margin-bottom:14px;display:flex;align-items:center;gap:7px;}
.syn-overview-label::before{content:'';display:inline-block;width:5px;height:5px;border-radius:50%;background:var(--gold);}
.syn-overview-text{font-family:var(--serif);font-size:clamp(.98rem,2.5vw,1.08rem);color:var(--t1);line-height:1.85;font-style:italic;}
.lit-discussion{margin-bottom:1.5rem;}
.lit-section-header{display:flex;align-items:center;gap:12px;margin-bottom:14px;padding-bottom:10px;border-bottom:1px solid var(--b1);}
.lit-section-num{font-family:var(--serif);font-size:1.4rem;color:var(--gold);font-style:italic;line-height:1;flex-shrink:0;width:28px;text-align:right;}
.lit-section-title{font-family:var(--sans);font-size:.98rem;font-weight:600;color:var(--t1);line-height:1.3;}
.lit-prose{font-family:var(--sans);font-size:.94rem;color:var(--t2);line-height:1.85;}
.syn-section{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:clamp(1.2rem,3.5vw,1.6rem) clamp(1.3rem,3.5vw,1.8rem);margin-bottom:1rem;transition:border-color .2s,box-shadow .2s;}
.syn-section:hover{border-color:var(--b2);box-shadow:0 6px 28px rgba(0,0,0,.22);}
.syn-head{font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:14px;display:flex;align-items:center;gap:9px;}
.syn-head-icon{width:26px;height:26px;border-radius:7px;background:var(--panel-2);border:1px solid var(--b1);display:flex;align-items:center;justify-content:center;font-size:.82rem;}
.syn-item{display:flex;gap:12px;margin-bottom:0;padding:10px 0;border-bottom:1px solid var(--b1);font-family:var(--sans);font-size:.93rem;color:var(--t2);line-height:1.65;}
.syn-item:last-child{border-bottom:none;padding-bottom:0;}
.syn-dot{width:6px;height:6px;border-radius:50%;flex-shrink:0;margin-top:9px;}
.dot-gold{background:var(--gold);}.dot-blue{background:var(--blue);}.dot-red{background:var(--red);}
.dot-green{background:var(--green);}.dot-violet{background:var(--violet);}
.var-tag{display:inline-block;background:var(--gold-soft);border:1px solid var(--gold-mid);color:var(--gold);font-family:var(--mono);font-size:.72rem;padding:4px 12px;border-radius:20px;margin:4px;}
.dom-method-card{display:flex;align-items:center;gap:16px;background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.2rem 1.5rem;margin-top:.75rem;transition:border-color .2s;}
.dom-method-card:hover{border-color:var(--b2);}
.dom-method-icon{width:44px;height:44px;border-radius:11px;background:var(--panel-3);border:1px solid var(--b2);display:flex;align-items:center;justify-content:center;font-size:1.3rem;flex-shrink:0;}
.dom-method-label{font-family:var(--mono);font-size:.6rem;letter-spacing:2px;text-transform:uppercase;color:var(--t3);margin-bottom:4px;}
.dom-method-val{font-family:var(--serif);font-size:1.05rem;color:var(--gold);font-style:italic;}

/* ── Export cards ── */
.exp-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.75rem 1.5rem;text-align:center;transition:all .2s;position:relative;overflow:hidden;cursor:default;}
.exp-card::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,var(--gold),var(--blue),transparent);opacity:0;transition:opacity .2s;}
.exp-card:hover{border-color:var(--b3);box-shadow:0 8px 32px rgba(0,0,0,.28);}
.exp-card:hover::before{opacity:1;}
.exp-icon{font-size:2.2rem;margin-bottom:12px;display:block;}
.exp-name{font-family:var(--sans);font-weight:600;font-size:1rem;color:var(--t1);margin-bottom:6px;}
.exp-desc{font-family:var(--sans);font-size:.8rem;color:var(--t3);line-height:1.55;}

/* ── Tables & tabs ── */
[data-testid="stDataFrame"]{border:1px solid var(--b1)!important;border-radius:var(--r-lg)!important;overflow:hidden!important;}
[data-testid="stDataFrame"] th{background:var(--panel)!important;color:var(--t3)!important;font-family:var(--mono)!important;font-size:.66rem!important;letter-spacing:1.5px!important;text-transform:uppercase!important;padding:11px 14px!important;border-bottom:1px solid var(--b1)!important;}
[data-testid="stDataFrame"] td{color:var(--t2)!important;font-family:var(--sans)!important;font-size:.88rem!important;padding:10px 14px!important;border-bottom:1px solid var(--b1)!important;}
[data-baseweb="tab-list"]{background:transparent!important;border-bottom:1px solid var(--b1)!important;gap:0!important;flex-wrap:wrap!important;}
[data-baseweb="tab"]{background:transparent!important;color:var(--t3)!important;font-family:var(--mono)!important;font-size:.68rem!important;letter-spacing:2px!important;text-transform:uppercase!important;padding:11px 20px!important;border:none!important;border-bottom:2px solid transparent!important;transition:all .15s!important;}
[data-baseweb="tab"]:hover{color:var(--t1)!important;}
[aria-selected="true"][data-baseweb="tab"]{color:var(--gold)!important;border-bottom-color:var(--gold)!important;}
.stProgress>div>div{background:linear-gradient(90deg,var(--gold),var(--gold-lt))!important;border-radius:4px!important;}
.stProgress>div{background:var(--panel)!important;border-radius:4px!important;height:4px!important;}
[data-testid="stExpander"]{background:var(--panel)!important;border:1px solid var(--b1)!important;border-radius:var(--r-lg)!important;}
[data-testid="stExpander"] summary{font-family:var(--mono)!important;font-size:.72rem!important;letter-spacing:1.5px!important;color:var(--t2)!important;padding:.8rem 1.1rem!important;}
[data-testid="stAlert"]{background:var(--panel)!important;border:1px solid var(--b2)!important;border-radius:var(--r)!important;color:var(--t2)!important;font-family:var(--sans)!important;font-size:.92rem!important;padding:.9rem 1.1rem!important;}
[data-baseweb="select"]{background:var(--bg-3)!important;border:1.5px solid var(--b2)!important;border-radius:var(--r)!important;}
[data-baseweb="select"]>div{background:var(--bg-3)!important;font-size:.93rem!important;font-family:var(--sans)!important;color:var(--t1)!important;min-height:2.9rem!important;border-radius:var(--r)!important;border-color:var(--b2)!important;}

/* ── Empty states ── */
.empty-st{text-align:center;padding:5rem 2rem;color:var(--t3);}
.empty-st-icon{font-size:3rem;opacity:.25;margin-bottom:18px;display:block;}
.empty-st-title{font-family:var(--serif);font-size:clamp(1.3rem,4vw,1.6rem);color:var(--t2);margin-bottom:10px;font-style:italic;}
.empty-st-desc{font-family:var(--sans);font-size:.92rem;line-height:1.7;max-width:360px;margin:0 auto;color:var(--t3);}

/* ── Processing cards ── */
.proc-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r);padding:13px 15px;margin-bottom:9px;display:flex;align-items:center;gap:12px;transition:all .3s ease;}
.proc-card.processing{border-left:3px solid var(--blue);}
.proc-card.success{border-left:3px solid var(--green);}
.proc-card.error{border-left:3px solid var(--red);}
.proc-title{font-family:var(--sans);font-weight:500;font-size:.88rem;color:var(--t1);margin-bottom:2px;}
.proc-sub{font-family:var(--mono);font-size:.78rem;color:var(--t3);}
.spinner{width:14px;height:14px;border:2px solid var(--b2);border-top:2px solid var(--blue);border-radius:50%;animation:spin .9s linear infinite;flex-shrink:0;}
@keyframes spin{0%{transform:rotate(0deg);}100%{transform:rotate(360deg);}}

/* ── Scrollbar ── */
::-webkit-scrollbar{width:4px;height:4px;}
::-webkit-scrollbar-track{background:var(--bg);}
::-webkit-scrollbar-thumb{background:var(--panel-3);border-radius:3px;}
::-webkit-scrollbar-thumb:hover{background:var(--gold);}

/* ── Animations ── */
@keyframes slideUp{from{opacity:0;transform:translateY(14px);}to{opacity:1;transform:translateY(0);}}
.anim-up{animation:slideUp .4s cubic-bezier(.16,1,.3,1) both;}
.anim-up-d1{animation-delay:.06s;}.anim-up-d2{animation-delay:.12s;}.anim-up-d3{animation-delay:.18s;}
hr{border-color:var(--b1)!important;margin:1.5rem 0!important;}

/* ── Mobile responsive ── */
@media(max-width:640px){
  [data-testid="stHorizontalBlock"]{flex-direction:column!important;gap:.75rem!important;}
  [data-testid="stHorizontalBlock"]>[data-testid="stVerticalBlock"]{width:100%!important;min-width:100%!important;flex:none!important;}
  .stButton>button{width:100%!important;}
  [data-testid="stMainBlockContainer"]{padding:1rem .9rem 3rem!important;}
  [data-baseweb="tab-list"]{overflow-x:auto!important;flex-wrap:nowrap!important;}
  .m-row{gap:8px!important;}.m-chip{min-width:80px!important;}
  .gate-card{padding:1.75rem 1.25rem!important;}
  .pricing-grid{grid-template-columns:1fr!important;}
  .gate-links{gap:12px!important;}
}
@media(max-width:480px){
  .ph-title{font-size:clamp(1.5rem,8vw,2rem)!important;}
  .gate-logo{font-size:clamp(1.7rem,8vw,2.3rem)!important;}
}
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════
# SVG LOGOS  — updated to a cleaner atom/lens icon
# ═══════════════════════════════════════════════════════════════════════
LOGO_SVG = """<svg width="36" height="36" viewBox="0 0 36 36" fill="none" xmlns="http://www.w3.org/2000/svg">
  <circle cx="18" cy="18" r="16" fill="#1b1b2f" stroke="#d4a843" stroke-width="1.4"/>
  <ellipse cx="18" cy="18" rx="7" ry="15" stroke="#d4a843" stroke-width="1.1" opacity="0.45" fill="none"/>
  <ellipse cx="18" cy="18" rx="15" ry="7" stroke="#d4a843" stroke-width="1.1" opacity="0.45" fill="none" transform="rotate(60 18 18)"/>
  <ellipse cx="18" cy="18" rx="15" ry="7" stroke="#d4a843" stroke-width="1.1" opacity="0.45" fill="none" transform="rotate(-60 18 18)"/>
  <circle cx="18" cy="18" r="3.5" fill="#f0c866"/>
  <circle cx="18" cy="18" r="1.6" fill="#fff" opacity="0.9"/>
</svg>"""

LOGO_SVG_GATE = """<svg width="56" height="56" viewBox="0 0 56 56" fill="none" xmlns="http://www.w3.org/2000/svg">
  <circle cx="28" cy="28" r="25" fill="#12122a" stroke="#d4a843" stroke-width="1.6"/>
  <ellipse cx="28" cy="28" rx="10" ry="23" stroke="#d4a843" stroke-width="1.3" opacity="0.5" fill="none"/>
  <ellipse cx="28" cy="28" rx="23" ry="10" stroke="#d4a843" stroke-width="1.3" opacity="0.5" fill="none" transform="rotate(60 28 28)"/>
  <ellipse cx="28" cy="28" rx="23" ry="10" stroke="#d4a843" stroke-width="1.3" opacity="0.5" fill="none" transform="rotate(-60 28 28)"/>
  <circle cx="28" cy="28" r="5.5" fill="#f0c866"/>
  <circle cx="28" cy="28" r="2.5" fill="#fff" opacity="0.9"/>
</svg>"""


# ═══════════════════════════════════════════════════════════════════════
# SESSION INIT
# ═══════════════════════════════════════════════════════════════════════
_DEFAULTS = {
    "authenticated": False, "page": "results",
    "queued_files": [], "extracted_papers": [],
    "synthesis_result": None, "processing_errors": [], "trigger_extract": False,
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ═══════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION UTILS
# ═══════════════════════════════════════════════════════════════════════
def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        chunks = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        doc.close()
        return "\n\n".join(chunks)
    except Exception as e:
        raise RuntimeError(f"PDF parsing failed: {e}")

def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        raise RuntimeError(f"DOCX parsing failed: {e}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):   return _extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"): return _extract_text_from_docx(file_bytes)
    elif name.endswith(".txt"):  return file_bytes.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {filename}")

def smart_truncate(text: str, max_tokens: int = 12000) -> str:
    max_chars = max_tokens * 4
    if len(text) <= max_chars: return text
    patterns = [r"(abstract[\s\S]{0,3000})", r"(introduction[\s\S]{0,2000})",
                r"(method(?:ology|s)?[\s\S]{0,4000})", r"(result[\s\S]{0,3000})",
                r"(discussion[\s\S]{0,2000})", r"(conclusion[\s\S]{0,2000})"]
    found = []
    text_lower = text.lower()
    for pat in patterns:
        m = re.search(pat, text_lower)
        if m: found.append(text[m.start():m.end()])
    return ("\n\n".join(found) if found else text)[:max_chars]

def format_file_size(size_bytes: int) -> str:
    if size_bytes < 1024: return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024: return f"{size_bytes/1024:.1f} KB"
    return f"{size_bytes/(1024*1024):.1f} MB"


# ═══════════════════════════════════════════════════════════════════════
# OPENAI HELPERS
# ═══════════════════════════════════════════════════════════════════════
EXTRACTION_PROMPT = """You are an expert academic research analyst. Extract structured empirical information from the research paper text below.

CRITICAL RULES:
1. Only extract information explicitly present in the text. Do NOT invent or hallucinate.
2. If a field cannot be determined, use "Not specified" — never guess.
3. Be concise but complete. Findings must be directional.
4. For methodology, be specific.

Return a single JSON object with EXACTLY these keys:
- author_year, title, research_context, methodology, independent_variables, dependent_variable,
  control_variables, findings, theoretical_contributions, practical_contributions,
  strengths, limitations, citation_apa, citation_mla, citation_harvard

Return ONLY the JSON object. No preamble, no markdown fences.

Paper text:
{text}"""

SYNTHESIS_PROMPT = """You are a senior academic researcher producing a literature review section. You have extractions from {n_papers} papers.

Produce a DEEP, CITATION-RICH synthesis as JSON with EXACTLY these keys:
overall_summary, discussion_convergence, discussion_conflicts, discussion_methodology,
discussion_gaps, discussion_theory, discussion_implications,
common_findings (list), conflicting_results (list), dominant_methodology (string),
methodology_patterns (list), common_weaknesses (list), research_gaps (list),
underexplored_variables (list), future_directions (list)

Each discussion_* key: 3-6 sentences of flowing academic prose with inline citations like (Smith et al., 2021).
Lists: concise strings with citations.

CRITICAL: Base ALL content on the actual extractions. Return ONLY valid JSON.

Paper extractions:
{extractions}"""

def _get_openai_client():
    api_key = os.environ.get("OPENAI_API_KEY","")
    if not api_key:
        raise ValueError("OpenAI API key not set. Add it in the sidebar Settings field.")
    try:
        from openai import OpenAI
        return OpenAI(api_key=api_key)
    except ImportError:
        raise RuntimeError("`openai` package not installed.")

def extract_paper(text: str, filename: str = "") -> dict:
    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role":"system","content":"Return valid JSON only."},
                  {"role":"user","content":EXTRACTION_PROMPT.format(text=text)}],
        temperature=0.1, max_tokens=1500, response_format={"type":"json_object"},
    )
    result = json.loads(response.choices[0].message.content)
    result["_source_file"] = filename; result["_status"] = "success"
    return result

def synthesize_papers(papers: list) -> dict:
    client = _get_openai_client()
    extractions_text = ""
    for i, p in enumerate(papers, 1):
        extractions_text += f"\n--- Paper {i} ---\n"
        for key in ["author_year","title","methodology","independent_variables","dependent_variable",
                    "control_variables","research_context","findings","theoretical_contributions","limitations","strengths"]:
            extractions_text += f"{key}: {p.get(key,'N/A')}\n"
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role":"system","content":"Return only valid JSON."},
                  {"role":"user","content":SYNTHESIS_PROMPT.format(n_papers=len(papers), extractions=extractions_text)}],
        temperature=0.25, max_tokens=4000, response_format={"type":"json_object"},
    )
    return json.loads(response.choices[0].message.content)


# ═══════════════════════════════════════════════════════════════════════
# EXPORTERS
# ═══════════════════════════════════════════════════════════════════════
DISPLAY_COLUMNS = [
    ("author_year","Author & Year"),("title","Title"),("research_context","Research Context"),
    ("methodology","Methodology"),("independent_variables","Independent Variables"),
    ("dependent_variable","Dependent Variable"),("control_variables","Control Variables"),
    ("findings","Key Findings"),("theoretical_contributions","Theoretical Contributions"),
    ("practical_contributions","Practical Contributions"),("strengths","Strengths"),("limitations","Limitations"),
]

def papers_to_csv(papers: list) -> bytes:
    output = io.StringIO()
    w = csv.writer(output)
    w.writerow([c[1] for c in DISPLAY_COLUMNS])
    for p in papers: w.writerow([p.get(c[0],"") for c in DISPLAY_COLUMNS])
    return output.getvalue().encode("utf-8")

def papers_to_excel(papers: list) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "EmpiricX Results"
    hf = Font(name="Calibri",bold=True,color="FFFFFF",size=10)
    hfill = PatternFill(start_color="07070d",end_color="07070d",fill_type="solid")
    ha = Alignment(horizontal="center",vertical="center",wrap_text=True)
    ca = Alignment(vertical="top",wrap_text=True)
    thin = Border(bottom=Side(style="thin",color="1b1b2f"),right=Side(style="thin",color="1b1b2f"))
    for ci,(key,label) in enumerate(DISPLAY_COLUMNS,1):
        cell = ws.cell(row=1,column=ci,value=label)
        cell.font=hf; cell.fill=hfill; cell.alignment=ha
    ws.row_dimensions[1].height = 36
    for ri,paper in enumerate(papers,2):
        for ci,(key,_) in enumerate(DISPLAY_COLUMNS,1):
            cell = ws.cell(row=ri,column=ci,value=paper.get(key,""))
            cell.alignment=ca; cell.border=thin; cell.font=Font(name="Calibri",size=9)
        fhex = "F5F5FA" if ri%2==0 else "FFFFFF"
        for ci in range(1,len(DISPLAY_COLUMNS)+1):
            ws.cell(row=ri,column=ci).fill = PatternFill(start_color=fhex,end_color=fhex,fill_type="solid")
    for ci,w in enumerate([18,30,22,18,22,18,18,35,30,30,22,25],1):
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.freeze_panes = "A2"
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

def synthesis_to_docx(synthesis: dict, papers: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document()
    for s in doc.sections:
        s.top_margin=Inches(1.1); s.bottom_margin=Inches(1.1)
        s.left_margin=Inches(1.3); s.right_margin=Inches(1.3)
    doc.styles["Normal"].font.name="Georgia"; doc.styles["Normal"].font.size=Pt(11)
    C_DARK=RGBColor(0x07,0x07,0x0d); C_GOLD=RGBColor(0xd4,0xa8,0x43); C_GOLD_D=RGBColor(0x9a,0x7a,0x2a)
    C_BLUE=RGBColor(0x55,0x85,0xff); C_MID=RGBColor(0x44,0x44,0x66); C_MUTED=RGBColor(0x88,0x88,0xaa)
    def shd_para(para,hx):
        pPr=para._p.get_or_add_pPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); pPr.append(shd)
    def shd_cell(cell,hx):
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)
    def add_rule(color="d4a843",weight="4"):
        rp=doc.add_paragraph(); pPr=rp._p.get_or_add_pPr(); pb=OxmlElement("w:pBdr")
        bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),weight)
        bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),color); pb.append(bot); pPr.append(pb)
        rp.paragraph_format.space_after=Pt(0)
    def add_sh(num,title,sub=""):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(2)
        rn=p.add_run(f"{num}.  "); rn.font.name="Georgia"; rn.font.size=Pt(14); rn.font.bold=True; rn.font.color.rgb=C_GOLD
        rt=p.add_run(title); rt.font.name="Georgia"; rt.font.size=Pt(14); rt.font.bold=True; rt.font.color.rgb=C_DARK
        if sub:
            s=doc.add_paragraph(sub); s.paragraph_format.space_before=Pt(0); s.paragraph_format.space_after=Pt(8)
            s.runs[0].font.name="Arial"; s.runs[0].font.size=Pt(9); s.runs[0].font.italic=True; s.runs[0].font.color.rgb=C_MUTED
        add_rule("d4a84333","2")
    def add_prose(text):
        para=doc.add_paragraph(); para.paragraph_format.space_before=Pt(6); para.paragraph_format.space_after=Pt(8)
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        run=para.add_run(text); run.font.name="Georgia"; run.font.size=Pt(11); run.font.color.rgb=C_MID
    def add_bullet(text):
        para=doc.add_paragraph(); para.paragraph_format.left_indent=Inches(0.3)
        para.paragraph_format.space_before=Pt(3); para.paragraph_format.space_after=Pt(3)
        rd=para.add_run("•  "); rd.font.name="Arial"; rd.font.size=Pt(10); rd.font.color.rgb=C_GOLD
        rt=para.add_run(str(text)); rt.font.name="Georgia"; rt.font.size=Pt(10.5); rt.font.color.rgb=C_MID
    # Cover
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    tr=tp.add_run("Empiri"); tr.font.name="Georgia"; tr.font.size=Pt(32); tr.font.bold=True; tr.font.color.rgb=C_DARK
    tx=tp.add_run("X"); tx.font.name="Georgia"; tx.font.size=Pt(32); tx.font.bold=True; tx.font.italic=True; tx.font.color.rgb=C_GOLD
    sub=doc.add_paragraph(); sub.add_run("Cross-Paper Literature Synthesis Report").font.color.rgb=C_MID
    meta=doc.add_paragraph()
    mr=meta.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}  ·  {len(papers)} paper(s) analysed")
    mr.font.name="Arial"; mr.font.size=Pt(9); mr.font.color.rgb=C_MUTED
    add_rule()
    summary=synthesis.get("overall_summary","")
    if summary:
        doc.add_paragraph(); lbl=doc.add_paragraph()
        lr=lbl.add_run("OVERVIEW"); lr.font.name="Arial"; lr.font.size=Pt(8); lr.font.bold=True; lr.font.color.rgb=C_GOLD
        sp=doc.add_paragraph(summary); sp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; sp.paragraph_format.space_after=Pt(12)
        shd_para(sp,"F0ECDE")
        if sp.runs: sp.runs[0].font.name="Georgia"; sp.runs[0].font.size=Pt(11.5); sp.runs[0].font.italic=True; sp.runs[0].font.color.rgb=RGBColor(0x33,0x22,0x00)
    doc.add_paragraph()
    for key,num,title,sub in [
        ("discussion_convergence","1","Convergence of Findings","Where the literature agrees."),
        ("discussion_conflicts","2","Conflicting Evidence","Divergences and contradictions."),
        ("discussion_methodology","3","Methodological Landscape","Critical appraisal of research designs."),
        ("discussion_theory","4","Theoretical Contributions","Frameworks invoked and theoretical advances."),
        ("discussion_gaps","5","Research Gaps","What remains unstudied or unresolved."),
        ("discussion_implications","6","Practical Implications","Actionable insights across all studies."),
    ]:
        prose=synthesis.get(key,"")
        if prose: add_sh(num,title,sub); add_prose(prose)
    doc.add_paragraph(); add_rule()
    qsh=doc.add_paragraph(); qr=qsh.add_run("Quick-Scan Summary")
    qr.font.name="Georgia"; qr.font.size=Pt(14); qr.font.bold=True; qr.font.color.rgb=C_DARK
    add_rule("5585ff","2")
    for key,title in [("common_findings","Common Findings"),("conflicting_results","Conflicting Results"),
                      ("methodology_patterns","Methodology Patterns"),("research_gaps","Research Gaps"),
                      ("common_weaknesses","Common Weaknesses"),("future_directions","Future Directions")]:
        items=synthesis.get(key,[])
        if items:
            sh=doc.add_paragraph(); sh.paragraph_format.space_before=Pt(10); sh.paragraph_format.space_after=Pt(4)
            sr2=sh.add_run(title); sr2.font.name="Arial"; sr2.font.size=Pt(10); sr2.font.bold=True; sr2.font.color.rgb=C_BLUE
            for item in items: add_bullet(item)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
    unexplored=synthesis.get("underexplored_variables",[])
    if unexplored:
        sh=doc.add_paragraph(); sh.paragraph_format.space_before=Pt(10)
        sr2=sh.add_run("Underexplored Variables"); sr2.font.name="Arial"; sr2.font.size=Pt(10); sr2.font.bold=True; sr2.font.color.rgb=C_BLUE
        vp=doc.add_paragraph(",  ".join(unexplored))
        if vp.runs: vp.runs[0].font.name="Georgia"; vp.runs[0].font.size=Pt(10.5); vp.runs[0].font.color.rgb=C_GOLD_D
    dom=synthesis.get("dominant_methodology","")
    if dom:
        sh=doc.add_paragraph(); sh.paragraph_format.space_before=Pt(10)
        sr2=sh.add_run("Dominant Methodology"); sr2.font.name="Arial"; sr2.font.size=Pt(10); sr2.font.bold=True; sr2.font.color.rgb=C_BLUE
        dp=doc.add_paragraph(dom); dp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        if dp.runs: dp.runs[0].font.name="Georgia"; dp.runs[0].font.size=Pt(11); dp.runs[0].font.italic=True; dp.runs[0].font.color.rgb=C_MID
    if papers:
        doc.add_page_break()
        ph=doc.add_paragraph(); ph.paragraph_format.space_after=Pt(10)
        pr=ph.add_run("Analysed Papers"); pr.font.name="Georgia"; pr.font.size=Pt(16); pr.font.bold=True; pr.font.color.rgb=C_DARK
        add_rule()
        col_w=[Cm(3.5),Cm(5.5),Cm(3),Cm(2.8),Cm(4.2)]
        headers=["Author & Year","Title","Methodology","Context","Key Finding (brief)"]
        table=doc.add_table(rows=1,cols=5); table.style="Table Grid"
        for i,w in enumerate(col_w): table.columns[i].width=w
        hdr_row=table.rows[0]
        for i,h_text in enumerate(headers):
            cell=hdr_row.cells[i]; cell.width=col_w[i]
            run=cell.paragraphs[0].add_run(h_text)
            run.font.bold=True; run.font.name="Arial"; run.font.size=Pt(8.5); run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            shd_cell(cell,"07070d")
        for idx,paper in enumerate(papers):
            ft=paper.get("findings","") or ""
            short=ft[:100]+("…" if len(ft)>100 else "")
            vals=[paper.get("author_year","—"),paper.get("title","—"),paper.get("methodology","—"),paper.get("research_context","—")[:60],short]
            fill="F0F0FA" if idx%2==0 else "FFFFFF"
            row=table.add_row()
            for i,val in enumerate(vals):
                cell=row.cells[i]; cell.width=col_w[i]
                run=cell.paragraphs[0].add_run(val)
                run.font.name="Arial"; run.font.size=Pt(8); run.font.color.rgb=C_MID
                shd_cell(cell,fill)
    doc.add_paragraph()
    fp=doc.add_paragraph()
    fr=fp.add_run("Generated by EmpiricX — AI-Powered Empirical Research Intelligence  ·  v2.0")
    fr.font.name="Arial"; fr.font.size=Pt(8); fr.font.italic=True; fr.font.color.rgb=C_MUTED
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
# API KEY BOOTSTRAP
# ═══════════════════════════════════════════════════════════════════════
try:
    key = st.secrets.get("OPENAI_API_KEY","")
    if key: os.environ["OPENAI_API_KEY"] = key
except Exception:
    pass


# ═══════════════════════════════════════════════════════════════════════
# GATE / LANDING PAGE
# ═══════════════════════════════════════════════════════════════════════
def _check_password() -> bool:
    try:    correct = st.secrets.get("APP_PASSWORD","empiricx2024")
    except: correct = "empiricx2024"
    if st.session_state.get("authenticated"): return True

    st.markdown('<div class="gate-bg"></div>', unsafe_allow_html=True)
    _, col_c, _ = st.columns([1, 1.8, 1])
    with col_c:
        st.markdown('<div class="gate-card">', unsafe_allow_html=True)

        # ── Logo + title
        st.markdown(f"""
        <div style="text-align:center;margin-bottom:1.5rem;">
          {LOGO_SVG_GATE}
          <div class="gate-logo" style="margin-top:12px;">Empiri<span>X</span></div>
          <div class="gate-sub">Research Intelligence Engine</div>
        </div>
        <div class="gate-pill">&#128270; AI-Powered &nbsp;·&nbsp; Empirical &nbsp;·&nbsp; Synthesis</div>
        <div class="gate-features">
          <div class="gate-feat"><div class="gate-feat-dot"></div>Extract structured empirical data from any research paper (PDF, DOCX, TXT)</div>
          <div class="gate-feat"><div class="gate-feat-dot"></div>Deep cross-paper synthesis — cited, flowing, publication-ready prose</div>
          <div class="gate-feat"><div class="gate-feat-dot"></div>Export-ready outputs: CSV, Excel &amp; full Word synthesis report</div>
        </div>""", unsafe_allow_html=True)

        # ── Pricing cards
        st.markdown("""
        <div class="pricing-section">
          <div class="pricing-section-title">Choose Your Plan</div>
          <div class="pricing-section-sub">Start free. Scale as your research grows.</div>
          <div class="pricing-grid">

            <div class="pricing-card">
              <div class="pricing-tier">Starter</div>
              <div class="pricing-price">Free<span> / mo</span></div>
              <div class="pricing-desc">Perfect for individual researchers and students.</div>
              <ul class="pricing-features">
                <li>Up to 3 papers per session</li>
                <li>CSV &amp; Excel export</li>
                <li>Basic extraction fields</li>
                <li>Community support</li>
              </ul>
              <a href="#" class="pricing-cta pricing-cta-outline" onclick="return false;">Get Started Free</a>
            </div>

            <div class="pricing-card pricing-card-popular">
              <div class="pricing-badge">Most Popular</div>
              <div class="pricing-tier">Pro</div>
              <div class="pricing-price">$19<span> / mo</span></div>
              <div class="pricing-desc">For serious researchers and research teams.</div>
              <ul class="pricing-features">
                <li>Unlimited papers per session</li>
                <li>Cross-paper synthesis</li>
                <li>Word report export (.docx)</li>
                <li>All citation formats</li>
                <li>Priority support</li>
              </ul>
              <a href="YOUR_FLUTTERWAVE_LINK_HERE" class="pricing-cta pricing-cta-primary" target="_blank" rel="noopener">Subscribe — $19/mo</a>
            </div>

            <div class="pricing-card">
              <div class="pricing-tier">Institutional</div>
              <div class="pricing-price">$79<span> / mo</span></div>
              <div class="pricing-desc">For labs, departments &amp; academic institutions.</div>
              <ul class="pricing-features">
                <li>Everything in Pro</li>
                <li>Up to 10 team seats</li>
                <li>Dedicated API access</li>
                <li>White-label export</li>
                <li>Onboarding &amp; SLA support</li>
              </ul>
              <a href="YOUR_FLUTTERWAVE_LINK_HERE" class="pricing-cta pricing-cta-outline" target="_blank" rel="noopener">Contact for Access</a>
            </div>

          </div>
        </div>""", unsafe_allow_html=True)

        # ── Login section
        st.markdown('<div style="border-top:1px solid var(--b1);padding-top:1.5rem;margin-top:.5rem"></div>', unsafe_allow_html=True)
        st.markdown('<div style="font-family:var(--mono);font-size:.6rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:.75rem;text-align:center">Already have access? Sign in below</div>', unsafe_allow_html=True)

        err_slot = st.empty()
        pwd = st.text_input("Access Password", type="password", placeholder="Enter your access password", key="pwd_input")
        enter = st.button("Enter Platform →", use_container_width=True)
        if enter or (pwd and pwd == correct):
            if pwd == correct:
                st.session_state["authenticated"] = True; st.rerun()
            elif pwd:
                err_slot.markdown('<div class="gate-error">⚠ Incorrect password — please try again.</div>', unsafe_allow_html=True)

        # ── Footer links: support + guide
        st.markdown("""
        <div class="gate-links">
          <a class="gate-link" href="mailto:support@empiricx.app">&#128233; Contact Support</a>
          <span class="gate-link-dot"></span>
          <a class="gate-link" href="https://docs.empiricx.app" target="_blank" rel="noopener">&#128218; User Guide</a>
          <span class="gate-link-dot"></span>
          <span class="gate-footer"><strong>EmpiricX</strong> v2.0 &nbsp;·&nbsp; Restricted access</span>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("</div>", unsafe_allow_html=True)
    return False

if not _check_password():
    st.stop()


# ═══════════════════════════════════════════════════════════════════════
# EXTRACTION RUNNER
# ═══════════════════════════════════════════════════════════════════════
def run_extraction_from_queue():
    queued = st.session_state.get("queued_files",[])
    extracted = st.session_state.get("extracted_papers",[])
    extracted_names = {p.get("_source_file") for p in extracted}
    pending = [f for f in queued if f["name"] not in extracted_names]
    if not pending: return
    total = len(pending); status = st.empty(); progress = st.progress(0.0)
    log_box = st.container(); errors = []
    status.info(f"🚀 Processing {total} paper(s)...")
    for i, finfo in enumerate(pending):
        fname = finfo["name"]; fobj = finfo.get("obj"); slot = log_box.empty(); t0 = time.time()
        slot.markdown(f'<div class="proc-card processing"><div class="spinner"></div><div><div class="proc-title">{fname}</div><div class="proc-sub">Parsing document...</div></div></div>', unsafe_allow_html=True)
        try:
            raw = fobj.read() if fobj else b""
            if not raw: raise ValueError("File is empty or could not be read.")
            text = extract_text(raw, fname)
            if len(text.strip()) < 150: raise ValueError("Too little text — file may be image-based.")
            text = smart_truncate(text, max_tokens=12000)
            slot.markdown(f'<div class="proc-card processing"><div class="spinner"></div><div><div class="proc-title">{fname}</div><div class="proc-sub">Extracting empirical insights...</div></div></div>', unsafe_allow_html=True)
            result = extract_paper(text, filename=fname)
            extracted.append(result)
            dur = round(time.time()-t0, 2)
            slot.markdown(f'<div class="proc-card success"><div>✔</div><div><div class="proc-title">{fname}</div><div class="proc-sub">Completed in {dur}s</div></div></div>', unsafe_allow_html=True)
        except Exception as e:
            errors.append(f"**{fname}**: {e}")
            slot.markdown(f'<div class="proc-card error"><div>⚠</div><div><div class="proc-title">{fname}</div><div class="proc-sub">{e}</div></div></div>', unsafe_allow_html=True)
        progress.progress((i+1)/total)
    st.session_state["extracted_papers"] = extracted
    st.session_state["synthesis_result"] = None
    st.session_state["processing_errors"] = errors
    progress.empty()
    if extracted: status.success(f"✅ {len(extracted)} paper(s) extracted and ready.")
    for e in errors: st.warning(e)


# ═══════════════════════════════════════════════════════════════════════
# PAGE: RESULTS
# ═══════════════════════════════════════════════════════════════════════
DISPLAY_COLS = [
    ("author_year","Author & Year"),("research_context","Research Context"),("methodology","Methodology"),
    ("independent_variables","IVs"),("dependent_variable","DV"),("findings","Key Findings"),
    ("strengths","Strengths"),("limitations","Limitations"),
]
FIELD_ICONS = {
    "research_context":"&#128205;","methodology":"&#9881;","independent_variables":"&#128204;",
    "dependent_variable":"&#127919;","control_variables":"&#128295;","findings":"&#128200;",
    "theoretical_contributions":"&#128161;","practical_contributions":"&#127981;",
    "strengths":"&#9989;","limitations":"&#9888;",
}

def _dblock(label, content):
    st.markdown(f'<div class="detail-block"><div class="detail-lbl">{label}</div><div class="detail-val">{content}</div></div>', unsafe_allow_html=True)

def page_results():
    papers = st.session_state.get("extracted_papers",[])
    st.markdown("""
    <div class="ph-wrap anim-up">
        <div class="ph-eye">Step 01 &middot; Analyse</div>
        <h1 class="ph-title">Extraction <span>Results</span></h1>
        <p class="ph-sub">Structured empirical data extracted from each paper. Inspect the full table or drill into any paper for detail.</p>
    </div>""", unsafe_allow_html=True)
    if not papers:
        st.markdown('<div class="empty-st"><span class="empty-st-icon">&#128230;</span><div class="empty-st-title">No papers extracted yet</div><div class="empty-st-desc">Upload papers using the sidebar, then click <strong>&#9889; Extract</strong> to begin.</div></div>', unsafe_allow_html=True)
        return
    methods = [p.get("methodology","") for p in papers if p.get("methodology") and p["methodology"]!="Not specified"]
    method_cnt = Counter(methods); top_m = method_cnt.most_common(1)[0][0][:24] if method_cnt else "—"
    dvs = [p.get("dependent_variable","") for p in papers if p.get("dependent_variable") and p["dependent_variable"]!="Not specified"]
    st.markdown(f"""
    <div class="m-row anim-up anim-up-d1">
        <div class="m-chip"><div class="m-val">{len(papers)}</div><div class="m-lbl">Papers Extracted</div></div>
        <div class="m-chip"><div class="m-val">{len(set(methods))}</div><div class="m-lbl">Unique Methods</div></div>
        <div class="m-chip"><div class="m-val">{len(set(dvs))}</div><div class="m-lbl">Outcome Variables</div></div>
        <div class="m-chip" style="min-width:180px"><div class="m-val" style="font-size:clamp(.9rem,2.2vw,1rem);padding-top:4px">{top_m}</div><div class="m-lbl">Top Methodology</div></div>
    </div>""", unsafe_allow_html=True)
    tab1, tab2 = st.tabs(["&#128202;  TABLE VIEW","&#128196;  PAPER DETAIL"])
    with tab1:
        rows = [{label:p.get(key,"—") for key,label in DISPLAY_COLS} for p in papers]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, height=min(620,90+65*len(papers)),
            column_config={"Author & Year":st.column_config.TextColumn(width="small"),"Research Context":st.column_config.TextColumn(width="medium"),
                           "Methodology":st.column_config.TextColumn(width="small"),"IVs":st.column_config.TextColumn(width="medium"),
                           "DV":st.column_config.TextColumn(width="small"),"Key Findings":st.column_config.TextColumn(width="large"),
                           "Strengths":st.column_config.TextColumn(width="medium"),"Limitations":st.column_config.TextColumn(width="medium")},
            hide_index=True)
    with tab2:
        options = [f"{i+1}. {p.get('author_year','Unknown')} — {p.get('_source_file','')}" for i,p in enumerate(papers)]
        idx = st.selectbox("Select paper", range(len(options)), format_func=lambda i: options[i], label_visibility="collapsed")
        p = papers[idx]
        st.markdown(f"""
        <div class="x-card x-card-gold anim-up" style="margin-bottom:1.25rem">
            <div style="font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--gold);margin-bottom:8px">{p.get("author_year","")}</div>
            <div style="font-family:var(--serif);font-size:1.2rem;color:var(--t1);margin-bottom:6px;font-style:italic;line-height:1.35">{p.get("title","Untitled")}</div>
            <div style="font-family:var(--mono);font-size:.68rem;color:var(--t3)">&#128196; {p.get("_source_file","")}</div>
        </div>""", unsafe_allow_html=True)
        labels = {"research_context":"Research Context","methodology":"Methodology","independent_variables":"Independent Variables",
                  "dependent_variable":"Dependent Variable","control_variables":"Control Variables","findings":"Key Findings",
                  "theoretical_contributions":"Theoretical Contributions","practical_contributions":"Practical Contributions",
                  "strengths":"Strengths","limitations":"Limitations"}
        c1,c2 = st.columns(2)
        with c1:
            for k in ["research_context","methodology","independent_variables","dependent_variable","control_variables"]:
                _dblock(f"{FIELD_ICONS.get(k,'')} {labels[k]}", p.get(k,"—"))
        with c2:
            for k in ["findings","theoretical_contributions","practical_contributions","strengths","limitations"]:
                _dblock(f"{FIELD_ICONS.get(k,'')} {labels[k]}", p.get(k,"—"))
        with st.expander("&#128218;  Citations — APA · MLA · Harvard"):
            for fmt,key in [("APA 7","citation_apa"),("MLA 9","citation_mla"),("Harvard","citation_harvard")]:
                st.markdown(f"**{fmt}**")
                st.code(p.get(key,"Not available"), language=None)
    st.markdown("---")
    # ── FIX #4: working routing CTAs
    c1, c2 = st.columns(2)
    with c1:
        if st.button("&#128279;  Run Cross-Paper Synthesis →", type="primary", use_container_width=True, key="results_go_synthesis"):
            st.session_state["page"] = "synthesis"
            st.rerun()
    with c2:
        if st.button("&#128229;  Export Data →", use_container_width=True, key="results_go_export"):
            st.session_state["page"] = "export"
            st.rerun()


# ═══════════════════════════════════════════════════════════════════════
# PAGE: SYNTHESIS
# ═══════════════════════════════════════════════════════════════════════
def _syn_block(title, items, dot_cls, subtitle=""):
    if not items: return
    items_html = "".join(f'<div class="syn-item"><div class="syn-dot {dot_cls}"></div><div>{item}</div></div>' for item in items)
    sub_html = f'<div style="font-family:var(--sans);font-size:.74rem;color:var(--t3);margin-bottom:10px">{subtitle}</div>' if subtitle else ""
    st.markdown(f'<div class="syn-section"><div class="syn-head"><div class="syn-head-icon">{title.split()[0]}</div>{" ".join(title.split()[1:])}</div>{sub_html}{items_html}</div>', unsafe_allow_html=True)

def page_synthesis():
    papers = st.session_state.get("extracted_papers",[]); synthesis = st.session_state.get("synthesis_result")
    st.markdown("""
    <div class="ph-wrap anim-up">
        <div class="ph-eye">Step 02 &middot; Synthesize</div>
        <h1 class="ph-title">Cross-Paper <span>Synthesis</span></h1>
        <p class="ph-sub">A deep, citation-anchored literature synthesis — convergences, conflicts, methodological critique, theoretical landscape, and research gaps — produced in flowing scholarly prose.</p>
    </div>""", unsafe_allow_html=True)
    if not papers:
        st.markdown('<div class="empty-st"><span class="empty-st-icon">&#128279;</span><div class="empty-st-title">No papers loaded</div><div class="empty-st-desc">Upload and extract at least 2 papers to run synthesis.</div></div>', unsafe_allow_html=True)
        return
    if len(papers) < 2: st.warning("Add at least 2 papers for meaningful cross-paper synthesis."); return
    c1,c2 = st.columns([4,1])
    with c1: st.markdown(f'<p style="font-family:var(--sans);color:var(--t2);font-size:.88rem;margin:0"><strong style="color:var(--gold)">{len(papers)}</strong> papers ready for deep synthesis.</p>', unsafe_allow_html=True)
    with c2: run = st.button("&#128279; Run" if not synthesis else "&#128260; Re-run", type="primary", use_container_width=True, key="syn_run_btn")
    if run:
        with st.spinner("Generating deep literature synthesis — this may take 30–60 seconds…"):
            try:
                result = synthesize_papers(papers); st.session_state["synthesis_result"]=result; synthesis=result
            except Exception as e: st.error(f"Synthesis failed: {e}"); return
    if not synthesis:
        st.markdown('<p style="font-family:var(--sans);color:var(--t3);font-size:.88rem;margin-top:1rem">Click <strong>Run</strong> above to generate the synthesis.</p>', unsafe_allow_html=True); return
    s = synthesis; st.markdown("---")
    summary = s.get("overall_summary","")
    if summary:
        st.markdown(f'<div class="syn-overview anim-up"><div class="syn-overview-label">Overview</div><div class="syn-overview-text">{summary}</div></div>', unsafe_allow_html=True)
    st.markdown('<div style="margin-bottom:1.25rem"><div style="font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:4px">&#128218; Literature Discussion</div><div style="font-family:var(--sans);font-size:.88rem;color:var(--t3)">Scholarly prose with inline citations — ready for your literature review.</div></div>', unsafe_allow_html=True)
    for key,num,title,subtitle in [
        ("discussion_convergence","1","Convergence of Findings","Where the literature agrees — patterns, consistencies, and shared insights."),
        ("discussion_conflicts","2","Conflicting Evidence","Divergences, contradictions, and contested findings."),
        ("discussion_methodology","3","Methodological Landscape","A critical appraisal of research designs and shared limitations."),
        ("discussion_theory","4","Theoretical Contributions","The frameworks invoked and advances this literature makes to theory."),
        ("discussion_gaps","5","Research Gaps","What remains unstudied, underexplored, or theoretically unresolved."),
        ("discussion_implications","6","Practical Implications","Actionable insights for managers and policymakers synthesised across all papers."),
    ]:
        prose = s.get(key,"")
        if prose:
            st.markdown(f"""
            <div class="syn-section lit-discussion anim-up">
                <div class="lit-section-header">
                    <div class="lit-section-num">{num}</div>
                    <div><div class="lit-section-title">{title}</div><div style="font-family:var(--sans);font-size:.78rem;color:var(--t3);margin-top:2px">{subtitle}</div></div>
                </div>
                <div class="lit-prose">{prose}</div>
            </div>""", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:1rem">&#9776; Quick-Scan Panels</div>', unsafe_allow_html=True)
    col1,col2 = st.columns(2)
    with col1:
        _syn_block("&#128200; Common Findings",    s.get("common_findings",[]),      "dot-gold",  "Recurring themes and results across papers")
        _syn_block("&#9881; Methodology Patterns",  s.get("methodology_patterns",[]), "dot-blue",  "Design choices and patterns observed")
        _syn_block("&#9888; Common Weaknesses",     s.get("common_weaknesses",[]),    "dot-red",   "Limitations shared across the literature")
    with col2:
        _syn_block("&#9889; Conflicting Results",   s.get("conflicting_results",[]),  "dot-red",   "Where studies disagree")
        _syn_block("&#128301; Research Gaps",       s.get("research_gaps",[]),        "dot-gold",  "Unstudied areas and open questions")
        _syn_block("&#128640; Future Directions",   s.get("future_directions",[]),    "dot-green", "Concrete recommendations for future research")
    unexplored = s.get("underexplored_variables",[])
    if unexplored:
        tags = "".join(f'<span class="var-tag">{v}</span>' for v in unexplored)
        st.markdown(f'<div class="syn-section" style="margin-top:.5rem"><div class="syn-head"><div class="syn-head-icon">&#128302;</div>Underexplored Variables</div><div style="margin-top:6px">{tags}</div></div>', unsafe_allow_html=True)
    dom = s.get("dominant_methodology","")
    if dom:
        st.markdown(f'<div class="dom-method-card"><div class="dom-method-icon">&#9881;</div><div><div class="dom-method-label">Dominant Methodology</div><div class="dom-method-val">{dom}</div></div></div>', unsafe_allow_html=True)
    st.markdown("---")
    # ── FIX #4: working export CTA on synthesis page
    if st.button("&#128229;  Export Synthesis Report →", type="primary", key="syn_go_export"):
        st.session_state["page"] = "export"
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════
# PAGE: EXPORT
# ═══════════════════════════════════════════════════════════════════════
def page_export():
    papers = st.session_state.get("extracted_papers",[]); synthesis = st.session_state.get("synthesis_result")
    st.markdown("""
    <div class="ph-wrap anim-up">
        <div class="ph-eye">Step 03 &middot; Export</div>
        <h1 class="ph-title">Export <span>Results</span></h1>
        <p class="ph-sub">Download your extracted data and the full literature synthesis report — polished and ready for your paper, thesis, or briefing.</p>
    </div>""", unsafe_allow_html=True)
    if not papers:
        st.markdown('<div class="empty-st"><span class="empty-st-icon">&#128229;</span><div class="empty-st-title">Nothing to export yet</div><div class="empty-st-desc">Extract papers first to enable downloads.</div></div>', unsafe_allow_html=True)
        return
    st.markdown(f'<div class="detail-lbl" style="margin-bottom:1rem">&#128202; Paper Extraction Data &nbsp;&middot;&nbsp; {len(papers)} paper(s)</div>', unsafe_allow_html=True)
    c1,c2 = st.columns(2)
    with c1:
        st.markdown('<div class="exp-card"><span class="exp-icon">&#128203;</span><div class="exp-name">CSV Spreadsheet</div><div class="exp-desc">Universal format &middot; Opens in any tool &middot; All fields included</div></div>', unsafe_allow_html=True)
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        try: st.download_button("&#8595;  Download CSV", data=papers_to_csv(papers), file_name="empiricx_results.csv", mime="text/csv", use_container_width=True, key="dl_csv")
        except Exception as e: st.error(f"CSV error: {e}")
    with c2:
        st.markdown('<div class="exp-card"><span class="exp-icon">&#128209;</span><div class="exp-name">Excel Workbook</div><div class="exp-desc">Styled .xlsx &middot; Freeze panes &middot; Alternating rows</div></div>', unsafe_allow_html=True)
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        try: st.download_button("&#8595;  Download Excel", data=papers_to_excel(papers), file_name="empiricx_results.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="dl_excel")
        except Exception as e: st.error(f"Excel error: {e}")
    st.markdown("---")
    st.markdown('<div class="detail-lbl" style="margin-bottom:1rem">&#128218; Synthesis Report &nbsp;&middot;&nbsp; Word Document</div>', unsafe_allow_html=True)
    if not synthesis:
        st.markdown('<div class="x-card" style="text-align:center;padding:2.5rem"><div style="font-size:2.5rem;margin-bottom:14px;opacity:.25">&#128196;</div><div style="font-family:var(--sans);font-size:.92rem;color:var(--t2);margin-bottom:18px;line-height:1.7">Run the Cross-Paper Synthesis first to generate the Word report.</div></div>', unsafe_allow_html=True)
        if st.button("&#128279;  Go to Synthesis →", key="export_go_synthesis"):
            st.session_state["page"] = "synthesis"
            st.rerun()
        return
    st.markdown('<div class="exp-card" style="text-align:left;max-width:520px"><div style="display:flex;align-items:center;gap:16px"><span style="font-size:2.4rem">&#128196;</span><div><div class="exp-name">Word Document (.docx)</div><div class="exp-desc" style="margin-top:6px">Full literature review report — scholarly overview, six deep discussion sections with inline citations, quick-scan summaries, and a reference table of all analysed papers.</div></div></div></div>', unsafe_allow_html=True)
    st.markdown('<div style="margin-top:12px"></div>', unsafe_allow_html=True)
    try: st.download_button("&#8595;  Download Synthesis Report (.docx)", data=synthesis_to_docx(synthesis,papers), file_name="empiricx_synthesis_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx")
    except Exception as e: st.error(f"DOCX error: {e}")


# ═══════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown(f'<div class="sb-brand"><div class="sb-logomark">{LOGO_SVG}</div><span class="sb-name">Empiri<em>X</em></span></div><div class="sb-tag">Research Intelligence</div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    # API key (if not set via secrets)
    if not os.environ.get("OPENAI_API_KEY"):
        st.markdown('<span class="sb-label">&#9881; OpenAI API Key</span>', unsafe_allow_html=True)
        api_key_input = st.text_input("API Key", type="password", placeholder="sk-...", label_visibility="collapsed", key="api_key_input")
        if api_key_input:
            os.environ["OPENAI_API_KEY"] = api_key_input
            st.success("Key saved for this session.")
        st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    st.markdown('<span class="sb-label">&#128196; Upload Papers</span>', unsafe_allow_html=True)
    uploaded = st.file_uploader("", type=["pdf","docx","txt"], accept_multiple_files=True, label_visibility="collapsed", key="sidebar_uploader")
    if uploaded:
        queued = st.session_state.get("queued_files",[]); existing = {f["name"] for f in queued}
        for f in uploaded:
            if f.name not in existing:
                queued.append({"name":f.name,"size":f.size,"obj":f}); existing.add(f.name)
        st.session_state["queued_files"] = queued

    queued = st.session_state.get("queued_files",[]); extracted = st.session_state.get("extracted_papers",[])
    extracted_names = {p.get("_source_file") for p in extracted}

    if queued:
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        for finfo in queued:
            _badge_cls = "sb-badge-ok" if finfo["name"] in extracted_names else "sb-badge-new"
            _badge_txt = "Done" if finfo["name"] in extracted_names else "Queued"
            st.markdown(f'<div class="sb-paper"><div class="sb-paper-icon">&#128196;</div><div style="min-width:0;flex:1"><div class="sb-paper-name">{finfo["name"]}</div><div class="sb-paper-meta">{format_file_size(finfo["size"])}</div></div><span class="sb-badge {_badge_cls}">{_badge_txt}</span></div>', unsafe_allow_html=True)
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        pending = [f for f in queued if f["name"] not in extracted_names]
        if pending:
            if st.button(f"⚡  Extract {len(pending)} paper{'s' if len(pending)>1 else ''}", use_container_width=True):
                st.session_state["trigger_extract"]=True; st.session_state["page"]="results"; st.rerun()
        if st.button("&#128465;  Clear all", use_container_width=True):
            st.session_state["queued_files"]=[]; st.session_state["extracted_papers"]=[]; st.session_state["synthesis_result"]=None; st.rerun()

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    st.markdown('<span class="sb-label">Navigation</span>', unsafe_allow_html=True)
    nav_map = {"&#128202;  Results":"results","&#128279;  Synthesis":"synthesis","&#128229;  Export":"export"}
    selected = st.radio("nav", list(nav_map.keys()),
                        index=list(nav_map.values()).index(st.session_state.get("page","results") if st.session_state.get("page") in nav_map.values() else "results"),
                        label_visibility="collapsed", key="main_nav")
    st.session_state["page"] = nav_map[selected]
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    n_papers = len(st.session_state.get("extracted_papers",[])); syn_done = "✓" if st.session_state.get("synthesis_result") else "—"
    st.markdown(f'<div class="sb-stats"><div class="sb-stat"><div class="sb-stat-n">{n_papers}</div><div class="sb-stat-l">Papers</div></div><div class="sb-stat"><div class="sb-stat-n">{syn_done}</div><div class="sb-stat-l">Synthesis</div></div></div>', unsafe_allow_html=True)
    st.markdown('<div style="margin-top:auto;padding-top:1rem"></div>', unsafe_allow_html=True)
    if st.button("&#9099;  Sign out", use_container_width=True):
        st.session_state["authenticated"]=False; st.rerun()


# ═══════════════════════════════════════════════════════════════════════
# ROUTER
# ═══════════════════════════════════════════════════════════════════════
if st.session_state.get("trigger_extract"):
    st.session_state["trigger_extract"] = False
    run_extraction_from_queue()

page = st.session_state.get("page","results")
if page == "results":    page_results()
elif page == "synthesis": page_synthesis()
elif page == "export":    page_export()
