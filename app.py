"""
EmpiricX — Research Intelligence Engine
Single-file Streamlit app. Deploy directly to Streamlit Community Cloud.

Requirements (requirements.txt):
    streamlit>=1.32.0
    openai>=1.30.0
    PyMuPDF>=1.24.0
    python-docx>=1.1.0
    openpyxl>=3.1.2
    pandas>=2.0.0
"""

# ═══════════════════════════════════════════════════════════════════════
# IMPORTS
# ═══════════════════════════════════════════════════════════════════════
import io
import os
import re
import csv
import json
import time
import datetime
import pandas as pd
import streamlit as st
from collections import Counter
from typing import Any, Optional


# ═══════════════════════════════════════════════════════════════════════
# GOOGLE SHEETS CREDIT SYSTEM
# ═══════════════════════════════════════════════════════════════════════
from google.oauth2.service_account import Credentials
import gspread

GSHEET_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

def _get_sheet():
    """Authenticate and return the EmpiricX credits worksheet."""
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=GSHEET_SCOPES)
    gc = gspread.authorize(creds)
    sh = gc.open_by_key(st.secrets["SHEET_ID"])
    return sh.sheet1

def lookup_key(access_key: str) -> dict | None:
    """
    Find a row by access key.
    Returns dict: row_index, key, credits, date_purchased, email  — or None.
    """
    try:
        ws = _get_sheet()
        records = ws.get_all_records()
        for i, row in enumerate(records, start=2):
            if str(row.get("Key", "")).strip() == access_key.strip():
                return {
                    "row_index": i,
                    "key": row["Key"],
                    "credits": int(row.get("Credits", 0)),
                    "date_purchased": row.get("DatePurchased", ""),
                    "email": row.get("Email", ""),
                }
        return None
    except Exception as e:
        st.error(f"Sheet lookup error: {e}")
        return None

def deduct_credits(row_index: int, current_credits: int, amount: int = 1) -> int:
    """Subtract `amount` credits and write back. Returns new credit count."""
    try:
        ws = _get_sheet()
        new_credits = max(0, current_credits - amount)
        ws.update_cell(row_index, 2, new_credits)   # Column B = Credits
        return new_credits
    except Exception as e:
        st.error(f"Credit deduction error: {e}")
        return current_credits

def synthesis_credit_cost(n_papers: int) -> int:
    """1 credit per 5 papers, rounded up."""
    import math
    return math.ceil(n_papers / 5)


# ═══════════════════════════════════════════════════════════════════════
# PAGE CONFIG  (must be first Streamlit call)
# ═══════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="EmpiricX — Research Intelligence",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ═══════════════════════════════════════════════════════════════════════
# GLOBAL CSS  (inlined from assets/style.css)
# ═══════════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Lora:ital,wght@0,400;0,600;0,700;1,400;1,600&family=Outfit:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --bg:#07070d;--bg-2:#0c0c18;--bg-3:#111120;
  --panel:#151526;--panel-2:#1b1b2f;--panel-3:#20203a;
  --gold:#d4a843;--gold-lt:#f0c866;--gold-dk:#9a7a2a;
  --gold-glow:rgba(212,168,67,.18);--gold-soft:rgba(212,168,67,.08);--gold-mid:rgba(212,168,67,.20);
  --blue:#5585ff;--blue-soft:rgba(85,133,255,.1);
  --green:#3ecf8e;--green-soft:rgba(62,207,142,.1);
  --red:#f05959;--red-soft:rgba(240,89,89,.1);
  --violet:#9b7eff;--violet-soft:rgba(155,126,255,.1);
  --t1:#f0f0f8;--t2:#9898b8;--t3:#52526a;--t4:#36364a;
  --b1:rgba(255,255,255,.055);--b2:rgba(255,255,255,.10);--b3:rgba(255,255,255,.16);
  --serif:'Lora',Georgia,serif;
  --sans:'Outfit',system-ui,sans-serif;
  --mono:'JetBrains Mono',monospace;
  --r:10px;--r-lg:16px;--r-xl:22px;
  --t-xs:clamp(.68rem,1.6vw,.78rem);--t-sm:clamp(.84rem,2vw,.95rem);
  --t-base:clamp(.96rem,2.4vw,1.05rem);--t-md:clamp(1.1rem,2.8vw,1.2rem);
  --t-lg:clamp(1.35rem,3.5vw,1.65rem);--t-xl:clamp(1.8rem,5vw,2.5rem);
  --t-2xl:clamp(2.2rem,6.5vw,3.1rem);
}

*,*::before,*::after{box-sizing:border-box;}
html{font-size:16px;-webkit-text-size-adjust:100%;}
html,body,[data-testid="stAppViewContainer"],[data-testid="stApp"]{
  background:var(--bg)!important;color:var(--t1)!important;
  font-family:var(--sans)!important;font-size:var(--t-base)!important;line-height:1.65!important;}
footer,[data-testid="stDecoration"]{display:none!important;}
[data-testid="stHeader"]{background:rgba(7,7,13,.9)!important;backdrop-filter:blur(12px)!important;border-bottom:1px solid var(--b1)!important;}
[data-testid="stHeader"] button{color:var(--t3)!important;background:transparent!important;}
[data-testid="stHeader"] button:hover{color:var(--gold)!important;background:var(--gold-soft)!important;}
[data-testid="stMainBlockContainer"]{padding:1.75rem clamp(1rem,4.5vw,3rem) 4rem!important;max-width:100%!important;}
@media(min-width:1200px){[data-testid="stMainBlockContainer"]{padding:2.25rem clamp(2rem,5.5vw,4.5rem) 5rem!important;}}
[data-testid="stSidebar"]{background:var(--bg-2)!important;border-right:1px solid var(--b1)!important;}
[data-testid="stSidebar"]>div:first-child{padding:1.5rem 1.25rem 2rem!important;}

.sb-brand{display:flex;align-items:center;gap:12px;padding:.2rem 0 .6rem;}
.sb-logomark{width:36px;height:36px;flex-shrink:0;display:block;}
.sb-name{font-family:var(--serif);font-size:1.35rem;font-weight:700;color:var(--t1);letter-spacing:-.02em;}
.sb-name em{color:var(--gold);font-style:italic;}
.sb-tag{font-family:var(--mono);font-size:.56rem;letter-spacing:3px;text-transform:uppercase;color:var(--t3);margin-bottom:1.1rem;padding-left:48px;opacity:.7;}
.sb-divider{height:1px;background:linear-gradient(90deg,var(--b2),transparent);margin:.85rem 0;}
.sb-label{font-family:var(--mono)!important;font-size:.58rem!important;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3)!important;margin-bottom:10px!important;display:block;opacity:.8;}
[data-testid="stSidebar"] [data-testid="stFileUploader"]{background:var(--bg-3)!important;border:1.5px dashed var(--b2)!important;border-radius:var(--r-lg)!important;transition:border-color .2s!important;}
[data-testid="stSidebar"] [data-testid="stFileUploader"]:hover{border-color:var(--gold)!important;}
[data-testid="stSidebar"] [data-testid="stFileUploader"] section{padding:1.1rem 1rem!important;}
[data-testid="stSidebar"] [data-testid="stFileUploader"] label,[data-testid="stSidebar"] [data-testid="stFileUploader"] p,[data-testid="stSidebar"] [data-testid="stFileUploader"] span{color:var(--t2)!important;font-size:.82rem!important;font-family:var(--sans)!important;}
/* Fix: hide the hidden label that causes "uploadUpload" duplication */
[data-testid="stSidebar"] [data-testid="stFileUploader"] [data-testid="stWidgetLabel"]{display:none!important;}
/* Ensure all uploader button text is clean */
[data-testid="stFileUploader"] button span{font-family:var(--sans)!important;font-size:.84rem!important;font-weight:500!important;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"]{background:var(--bg-3)!important;border:1.5px dashed var(--b2)!important;border-radius:var(--r-lg)!important;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"]:hover{border-color:var(--gold-dk)!important;}
.sb-paper{display:flex;align-items:flex-start;gap:9px;background:var(--bg-3);border:1px solid var(--b1);border-radius:var(--r);padding:10px 12px;margin-bottom:6px;transition:border-color .15s;}
.sb-paper:hover{border-color:var(--b2);}
.sb-paper-icon{width:28px;height:28px;border-radius:7px;background:var(--panel-2);border:1px solid var(--b1);display:flex;align-items:center;justify-content:center;font-size:.9rem;flex-shrink:0;margin-top:1px;}
.sb-paper-name{font-size:.82rem;font-weight:500;color:var(--t1);line-height:1.3;word-break:break-word;}
.sb-paper-meta{font-family:var(--mono);font-size:.64rem;color:var(--t3);margin-top:3px;}
.sb-badge{margin-left:auto;flex-shrink:0;font-family:var(--mono);font-size:.6rem;letter-spacing:.5px;padding:3px 9px;border-radius:20px;white-space:nowrap;}
.sb-badge-ok{background:var(--green-soft);color:var(--green);border:1px solid rgba(62,207,142,.25);}
.sb-badge-new{background:var(--blue-soft);color:var(--blue);border:1px solid rgba(85,133,255,.25);}
.sb-stats{display:flex;gap:8px;margin:.6rem 0;}
.sb-stat{flex:1;background:var(--bg-3);border:1px solid var(--b1);border-radius:var(--r);padding:14px 8px;text-align:center;transition:border-color .2s;}
.sb-stat:hover{border-color:var(--b2);}
.sb-stat-n{font-family:var(--serif);font-size:1.75rem;color:var(--gold);line-height:1;font-weight:700;}
.sb-stat-l{font-family:var(--mono);font-size:.55rem;letter-spacing:1.8px;text-transform:uppercase;color:var(--t3);margin-top:5px;}
[data-testid="stRadio"] label{font-family:var(--sans)!important;font-size:.9rem!important;font-weight:400!important;color:var(--t2)!important;padding:10px 14px!important;border-radius:var(--r)!important;border:1px solid transparent!important;transition:all .15s!important;cursor:pointer!important;line-height:1.45!important;letter-spacing:.01em!important;}
[data-testid="stRadio"] label:hover{color:var(--t1)!important;background:var(--panel)!important;border-color:var(--b1)!important;}
[data-testid="stRadio"] [aria-checked="true"] label{color:var(--gold)!important;background:var(--gold-soft)!important;border-color:var(--gold-mid)!important;font-weight:500!important;}
.stButton>button{background:linear-gradient(135deg,#d4a843 0%,#a87e2a 100%)!important;color:#0a0700!important;border:none!important;border-radius:var(--r)!important;font-family:var(--sans)!important;font-weight:600!important;font-size:.9rem!important;letter-spacing:.2px!important;padding:.65rem 1.6rem!important;min-height:2.65rem!important;transition:all .2s cubic-bezier(.16,1,.3,1)!important;box-shadow:0 2px 12px var(--gold-glow)!important;}
/* Secondary (ghost) button — Sign Out, Clear, utility actions */
[data-testid="stSidebar"] .stButton:not(:has([key="sidebar_extract"])) button,
.stButton[data-testid*="signout"] button,
.stButton[data-testid*="clear"] button{background:transparent!important;border:1px solid var(--b2)!important;color:var(--t3)!important;box-shadow:none!important;font-weight:400!important;}
[data-testid="stSidebar"] .stButton:not(:has([key="sidebar_extract"])) button:hover{border-color:var(--b3)!important;color:var(--t2)!important;background:var(--panel)!important;transform:none!important;box-shadow:none!important;}
.stButton>button:hover{transform:translateY(-2px)!important;box-shadow:0 8px 28px rgba(212,168,67,.32)!important;background:linear-gradient(135deg,#e8bc55 0%,#bc9035 100%)!important;}
.stButton>button:active{transform:translateY(0) scale(.98)!important;}
/* Sidebar utility/ghost buttons */
button[kind="secondary"]{background:transparent!important;border:1px solid var(--b2)!important;color:var(--t2)!important;box-shadow:none!important;}
button[kind="secondary"]:hover{border-color:var(--gold)!important;color:var(--gold)!important;background:var(--gold-soft)!important;transform:none!important;box-shadow:none!important;}
[data-testid="stBaseButton-secondary"]>button{background:transparent!important;border:1px solid var(--b2)!important;color:var(--t2)!important;box-shadow:none!important;}
[data-testid="stBaseButton-secondary"]>button:hover{border-color:var(--gold)!important;color:var(--gold)!important;background:var(--gold-soft)!important;transform:none!important;box-shadow:none!important;}
[data-testid="stDownloadButton"]>button{min-height:2.65rem!important;font-size:.9rem!important;}
.stTextInput input,[data-baseweb="input"] input{background:var(--bg-3)!important;border:1.5px solid var(--b2)!important;border-radius:var(--r)!important;color:var(--t1)!important;font-family:var(--sans)!important;font-size:1rem!important;padding:.7rem 1.1rem!important;min-height:2.9rem!important;transition:border-color .2s,box-shadow .2s!important;}
.stTextInput input:focus,[data-baseweb="input"] input:focus{border-color:var(--gold)!important;box-shadow:0 0 0 3px var(--gold-glow)!important;outline:none!important;}
.stTextInput label{font-family:var(--mono)!important;font-size:.66rem!important;letter-spacing:2.5px!important;text-transform:uppercase!important;color:var(--t3)!important;margin-bottom:8px!important;}
.ph-wrap{padding:1.5rem 0 1.6rem;border-bottom:1px solid var(--b1);margin-bottom:1.75rem;position:relative;}
.ph-wrap::after{content:'';position:absolute;bottom:-1px;left:0;width:64px;height:2px;background:linear-gradient(90deg,var(--gold),transparent);}
.ph-eye{display:inline-flex;align-items:center;gap:7px;font-family:var(--mono);font-size:.66rem;letter-spacing:3px;text-transform:uppercase;color:var(--gold);margin-bottom:11px;background:var(--gold-soft);border:1px solid var(--gold-mid);border-radius:20px;padding:4px 12px 4px 10px;}
.ph-eye::before{content:'\25C6';font-size:.45rem;opacity:.7;}
.ph-title{font-family:var(--serif);font-size:var(--t-xl);color:var(--t1);line-height:1.1;margin:0;font-weight:700;}
.ph-title span{color:var(--gold);font-style:italic;}
.ph-sub{font-size:.97rem;color:var(--t2);margin-top:12px;max-width:580px;line-height:1.7;}
.gate-bg{position:fixed;inset:0;z-index:0;background:radial-gradient(ellipse 80% 60% at 15% 110%,rgba(212,168,67,.06) 0%,transparent 60%),radial-gradient(ellipse 60% 50% at 88% -5%,rgba(85,133,255,.05) 0%,transparent 55%),var(--bg);pointer-events:none;}
.gate-bg::before{content:'';position:absolute;inset:0;background-image:linear-gradient(rgba(255,255,255,.018) 1px,transparent 1px),linear-gradient(90deg,rgba(255,255,255,.018) 1px,transparent 1px);background-size:56px 56px;mask-image:radial-gradient(ellipse 75% 75% at 50% 50%,black,transparent);animation:gridDrift 25s linear infinite;}
@keyframes gridDrift{0%{transform:translate(0,0);}100%{transform:translate(56px,56px);}}
.gate-card{position:relative;z-index:1;background:rgba(18,18,30,.92);border:1px solid var(--b2);border-radius:var(--r-xl);padding:clamp(2rem,6vw,3.25rem) clamp(1.75rem,6vw,3.75rem);width:100%;backdrop-filter:blur(24px);box-shadow:0 40px 80px rgba(0,0,0,.55),0 0 0 1px rgba(212,168,67,.05);animation:cardRise .55s cubic-bezier(.16,1,.3,1) both;}
@keyframes cardRise{from{opacity:0;transform:translateY(24px) scale(.97);}to{opacity:1;transform:translateY(0) scale(1);}}
.gate-logo{text-align:center;margin-bottom:.35rem;font-family:var(--serif);font-size:clamp(2rem,7.5vw,2.8rem);font-weight:700;color:var(--t1);}
.gate-logo span{color:var(--gold);font-style:italic;}
.gate-sub{text-align:center;margin-bottom:clamp(1rem,3vw,1.5rem);font-family:var(--mono);font-size:.64rem;letter-spacing:3px;text-transform:uppercase;color:var(--t3);}
.gate-pill{display:flex;align-items:center;justify-content:center;gap:7px;background:var(--gold-soft);border:1px solid var(--gold-mid);border-radius:20px;padding:6px 16px;font-family:var(--mono);font-size:.64rem;color:var(--gold);letter-spacing:1px;margin:0 auto clamp(1.5rem,4vw,2rem);width:fit-content;}
.gate-features{display:flex;flex-direction:column;gap:10px;margin-bottom:clamp(1.5rem,4vw,2.25rem);}
.gate-feat{display:flex;align-items:flex-start;gap:11px;font-size:.9rem;color:var(--t2);line-height:1.5;}
.gate-feat-dot{width:6px;height:6px;border-radius:50%;background:var(--gold);flex-shrink:0;margin-top:8px;}
.gate-error{background:var(--red-soft);border:1px solid rgba(240,89,89,.28);border-radius:var(--r);padding:11px 16px;font-size:.88rem;color:var(--red);margin-bottom:1rem;text-align:center;}
.gate-footer{text-align:center;margin-top:clamp(1rem,3vw,1.75rem);font-size:.74rem;color:var(--t3);}
.gate-footer strong{color:var(--gold);font-weight:500;}
.m-row{display:flex;gap:14px;flex-wrap:wrap;margin-bottom:2rem;}
.m-chip{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.2rem 1.5rem;flex:1;min-width:120px;position:relative;overflow:hidden;transition:border-color .2s,box-shadow .2s;}
.m-chip::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,var(--gold),transparent);opacity:.5;}
.m-chip:hover{border-color:var(--b2);box-shadow:0 4px 20px rgba(0,0,0,.28);}
.m-val{font-family:var(--serif);font-size:clamp(1.6rem,4.5vw,2.1rem);color:var(--gold);line-height:1;margin-bottom:5px;}
.m-lbl{font-family:var(--mono);font-size:.62rem;letter-spacing:1.5px;text-transform:uppercase;color:var(--t3);}
.x-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:clamp(1.1rem,3vw,1.5rem);margin-bottom:1rem;transition:border-color .2s,box-shadow .2s;}
.x-card:hover{border-color:var(--b2);box-shadow:0 4px 24px rgba(0,0,0,.25);}
.x-card-gold{border-left:3px solid var(--gold);}
.detail-block{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r);padding:1rem 1.2rem;margin-bottom:.75rem;transition:border-color .15s;}
.detail-block:hover{border-color:var(--b2);}
.detail-lbl{font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:7px;}
.detail-val{font-size:.93rem;color:var(--t2);line-height:1.65;}
.syn-overview{background:linear-gradient(135deg,var(--panel) 0%,var(--panel-2) 100%);border:1px solid var(--b2);border-left:4px solid var(--gold);border-radius:var(--r-lg);padding:clamp(1.4rem,4vw,2rem) clamp(1.5rem,4vw,2.2rem);margin-bottom:1.75rem;position:relative;}
.syn-overview-label{font-family:var(--mono);font-size:.6rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--gold);margin-bottom:14px;display:flex;align-items:center;gap:7px;}
.syn-overview-label::before{content:'';display:inline-block;width:5px;height:5px;border-radius:50%;background:var(--gold);}
.syn-overview-text{font-family:var(--serif);font-size:clamp(.98rem,2.5vw,1.08rem);color:var(--t1);line-height:1.85;font-style:italic;}
.lit-discussion{margin-bottom:1.5rem;}
.lit-section-header{display:flex;align-items:center;gap:12px;margin-bottom:14px;padding-bottom:10px;border-bottom:1px solid var(--b1);}
.lit-section-num{font-family:var(--serif);font-size:1.4rem;color:var(--gold);font-style:italic;line-height:1;flex-shrink:0;width:28px;text-align:right;}
.lit-section-title{font-family:var(--sans);font-size:.98rem;font-weight:600;color:var(--t1);line-height:1.3;}
.lit-prose{font-size:.94rem;color:var(--t2);line-height:1.85;}
.syn-section{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:clamp(1.2rem,3.5vw,1.6rem) clamp(1.3rem,3.5vw,1.8rem);margin-bottom:1rem;transition:border-color .2s,box-shadow .2s;}
.syn-section:hover{border-color:var(--b2);box-shadow:0 6px 28px rgba(0,0,0,.22);}
.syn-head{font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:14px;display:flex;align-items:center;gap:9px;}
.syn-head-icon{width:26px;height:26px;border-radius:7px;background:var(--panel-2);border:1px solid var(--b1);display:flex;align-items:center;justify-content:center;font-size:.82rem;}
.syn-item{display:flex;gap:12px;margin-bottom:0;padding:10px 0;border-bottom:1px solid var(--b1);font-size:.93rem;color:var(--t2);line-height:1.65;}
.syn-item:last-child{border-bottom:none;padding-bottom:0;}
.syn-dot{width:6px;height:6px;border-radius:50%;flex-shrink:0;margin-top:9px;}
.dot-gold{background:var(--gold);}.dot-blue{background:var(--blue);}.dot-red{background:var(--red);}
.dot-green{background:var(--green);}.dot-violet{background:var(--violet);}
.var-tag{display:inline-block;background:var(--gold-soft);border:1px solid var(--gold-mid);color:var(--gold);font-family:var(--mono);font-size:.72rem;padding:4px 12px;border-radius:20px;margin:4px;}
.dom-method-card{display:flex;align-items:center;gap:16px;background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.2rem 1.5rem;margin-top:.75rem;transition:border-color .2s;}
.dom-method-card:hover{border-color:var(--b2);}
.dom-method-icon{width:44px;height:44px;border-radius:11px;background:var(--panel-3);border:1px solid var(--b2);display:flex;align-items:center;justify-content:center;font-size:1.3rem;flex-shrink:0;}
.dom-method-label{font-family:var(--mono);font-size:.6rem;letter-spacing:2px;text-transform:uppercase;color:var(--t3);margin-bottom:4px;}
.dom-method-val{font-family:var(--serif);font-size:1.05rem;color:var(--gold);font-style:italic;}
.exp-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r-lg);padding:1.75rem 1.5rem;text-align:center;transition:all .2s;position:relative;overflow:hidden;cursor:default;}
.exp-card::before{content:'';position:absolute;top:0;left:0;right:0;height:2px;background:linear-gradient(90deg,var(--gold),var(--blue),transparent);opacity:0;transition:opacity .2s;}
.exp-card:hover{border-color:var(--b3);box-shadow:0 8px 32px rgba(0,0,0,.28);}
.exp-card:hover::before{opacity:1;}
.exp-icon{font-size:2.2rem;margin-bottom:12px;display:block;}
.exp-name{font-weight:600;font-size:1rem;color:var(--t1);margin-bottom:6px;}
.exp-desc{font-size:.8rem;color:var(--t3);line-height:1.55;}
[data-testid="stDataFrame"]{border:1px solid var(--b1)!important;border-radius:var(--r-lg)!important;overflow:hidden!important;}
[data-testid="stDataFrame"] th{background:var(--panel)!important;color:var(--t3)!important;font-family:var(--mono)!important;font-size:.66rem!important;letter-spacing:1.5px!important;text-transform:uppercase!important;padding:11px 14px!important;border-bottom:1px solid var(--b1)!important;}
[data-testid="stDataFrame"] td{color:var(--t2)!important;font-size:.88rem!important;padding:10px 14px!important;border-bottom:1px solid var(--b1)!important;}
[data-baseweb="tab-list"]{background:transparent!important;border-bottom:1px solid var(--b1)!important;gap:0!important;flex-wrap:wrap!important;}
[data-baseweb="tab"]{background:transparent!important;color:var(--t3)!important;font-family:var(--mono)!important;font-size:.68rem!important;letter-spacing:2px!important;text-transform:uppercase!important;padding:11px 20px!important;border:none!important;border-bottom:2px solid transparent!important;transition:all .15s!important;}
[data-baseweb="tab"]:hover{color:var(--t1)!important;}
[aria-selected="true"][data-baseweb="tab"]{color:var(--gold)!important;border-bottom-color:var(--gold)!important;}
.stProgress>div>div{background:linear-gradient(90deg,var(--gold),var(--gold-lt))!important;border-radius:4px!important;}
.stProgress>div{background:var(--panel)!important;border-radius:4px!important;height:4px!important;}
[data-testid="stExpander"]{background:var(--panel)!important;border:1px solid var(--b1)!important;border-radius:var(--r-lg)!important;}
[data-testid="stExpander"] summary{font-family:var(--mono)!important;font-size:.72rem!important;letter-spacing:1.5px!important;color:var(--t2)!important;padding:.8rem 1.1rem!important;}
[data-testid="stAlert"]{background:var(--panel)!important;border:1px solid var(--b2)!important;border-radius:var(--r)!important;color:var(--t2)!important;font-size:.92rem!important;padding:.9rem 1.1rem!important;}
[data-baseweb="select"]{background:var(--bg-3)!important;border:1.5px solid var(--b2)!important;border-radius:var(--r)!important;}
[data-baseweb="select"]>div{background:var(--bg-3)!important;font-size:.93rem!important;font-family:var(--sans)!important;color:var(--t1)!important;min-height:2.9rem!important;border-radius:var(--r)!important;border-color:var(--b2)!important;}
.empty-st{text-align:center;padding:5rem 2rem;color:var(--t3);}
.empty-st-icon{font-size:3rem;opacity:.25;margin-bottom:18px;display:block;}
.empty-st-title{font-family:var(--serif);font-size:clamp(1.3rem,4vw,1.6rem);color:var(--t2);margin-bottom:10px;font-style:italic;}
.empty-st-desc{font-size:.92rem;line-height:1.7;max-width:360px;margin:0 auto;color:var(--t3);}
.proc-card{background:var(--panel);border:1px solid var(--b1);border-radius:var(--r);padding:13px 15px;margin-bottom:9px;display:flex;align-items:center;gap:12px;transition:all .3s ease;}
.proc-card.processing{border-left:3px solid var(--blue);}
.proc-card.success{border-left:3px solid var(--green);}
.proc-card.error{border-left:3px solid var(--red);}
.proc-title{font-weight:500;font-size:.88rem;color:var(--t1);margin-bottom:2px;}
.proc-sub{font-size:.78rem;color:var(--t3);font-family:var(--mono);}
.spinner{width:14px;height:14px;border:2px solid var(--b2);border-top:2px solid var(--blue);border-radius:50%;animation:spin .9s linear infinite;flex-shrink:0;}
@keyframes spin{0%{transform:rotate(0deg);}100%{transform:rotate(360deg);}}
::-webkit-scrollbar{width:4px;height:4px;}
::-webkit-scrollbar-track{background:var(--bg);}
::-webkit-scrollbar-thumb{background:var(--panel-3);border-radius:3px;}
::-webkit-scrollbar-thumb:hover{background:var(--gold);}
@keyframes slideUp{from{opacity:0;transform:translateY(14px);}to{opacity:1;transform:translateY(0);}}
.anim-up{animation:slideUp .4s cubic-bezier(.16,1,.3,1) both;}
.anim-up-d1{animation-delay:.06s;}.anim-up-d2{animation-delay:.12s;}.anim-up-d3{animation-delay:.18s;}
hr{border-color:var(--b1)!important;margin:1.5rem 0!important;}
@media(max-width:640px){
  [data-testid="stHorizontalBlock"]{flex-direction:column!important;gap:.75rem!important;}
  [data-testid="stHorizontalBlock"]>[data-testid="stVerticalBlock"]{width:100%!important;min-width:100%!important;flex:none!important;}
  .stButton>button{width:100%!important;}
  [data-testid="stMainBlockContainer"]{padding:1rem .9rem 3rem!important;}
  [data-baseweb="tab-list"]{overflow-x:auto!important;flex-wrap:nowrap!important;}
  .m-row{gap:8px!important;}.m-chip{min-width:80px!important;}
}
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════
# SVG LOGOS
# ═══════════════════════════════════════════════════════════════════════
LOGO_SVG = """<svg width="36" height="36" viewBox="0 0 36 36" fill="none" xmlns="http://www.w3.org/2000/svg">
  <!-- Document stack base -->
  <rect x="5" y="7" width="19" height="24" rx="2.5" fill="#1b1b2f" stroke="#d4a843" stroke-width="1.2" opacity="0.6"/>
  <rect x="8" y="4" width="19" height="24" rx="2.5" fill="#20203a" stroke="#d4a843" stroke-width="1.2" opacity="0.8"/>
  <rect x="11" y="1" width="19" height="24" rx="2.5" fill="#1b1b2f" stroke="#d4a843" stroke-width="1.4"/>
  <!-- Document lines -->
  <line x1="15" y1="8" x2="26" y2="8" stroke="#d4a843" stroke-width="1" opacity="0.5"/>
  <line x1="15" y1="11.5" x2="26" y2="11.5" stroke="#d4a843" stroke-width="1" opacity="0.35"/>
  <line x1="15" y1="15" x2="22" y2="15" stroke="#d4a843" stroke-width="1" opacity="0.25"/>
  <!-- Magnifier lens -->
  <circle cx="24" cy="25" r="7.5" fill="#07070d" stroke="#d4a843" stroke-width="1.6"/>
  <circle cx="24" cy="25" r="4.5" fill="none" stroke="#f0c866" stroke-width="1" opacity="0.6"/>
  <circle cx="22.5" cy="23.5" r="1.2" fill="#f0c866" opacity="0.7"/>
  <!-- Handle -->
  <line x1="29.5" y1="30.5" x2="33" y2="34" stroke="#d4a843" stroke-width="2" stroke-linecap="round"/>
</svg>"""

LOGO_SVG_GATE = """<svg width="64" height="64" viewBox="0 0 64 64" fill="none" xmlns="http://www.w3.org/2000/svg">
  <!-- Outer glow ring -->
  <circle cx="32" cy="32" r="30" fill="none" stroke="#d4a843" stroke-width="0.5" opacity="0.2"/>
  <!-- Document stack -->
  <rect x="8" y="13" width="32" height="40" rx="4" fill="#12122a" stroke="#d4a843" stroke-width="1.2" opacity="0.5"/>
  <rect x="12" y="9" width="32" height="40" rx="4" fill="#1b1b2f" stroke="#d4a843" stroke-width="1.3" opacity="0.75"/>
  <rect x="16" y="5" width="32" height="40" rx="4" fill="#12122a" stroke="#d4a843" stroke-width="1.6"/>
  <!-- Document lines representing data -->
  <line x1="22" y1="14" x2="42" y2="14" stroke="#d4a843" stroke-width="1.4" opacity="0.7"/>
  <line x1="22" y1="19.5" x2="42" y2="19.5" stroke="#d4a843" stroke-width="1.1" opacity="0.45"/>
  <line x1="22" y1="25" x2="36" y2="25" stroke="#d4a843" stroke-width="1.1" opacity="0.3"/>
  <line x1="22" y1="30.5" x2="39" y2="30.5" stroke="#d4a843" stroke-width="1.1" opacity="0.25"/>
  <!-- Search/analysis lens -->
  <circle cx="41" cy="46" r="12" fill="#07070d" stroke="#d4a843" stroke-width="2"/>
  <circle cx="41" cy="46" r="7.5" fill="none" stroke="#f0c866" stroke-width="1.4" opacity="0.7"/>
  <!-- Inner shine -->
  <circle cx="38.5" cy="43.5" r="2" fill="#f0c866" opacity="0.55"/>
  <!-- Handle -->
  <line x1="50.5" y1="55.5" x2="57" y2="62" stroke="#d4a843" stroke-width="2.8" stroke-linecap="round"/>
  <!-- Node dots suggesting data intelligence -->
  <circle cx="22" cy="42" r="1.5" fill="#5585ff" opacity="0.7"/>
  <circle cx="28" cy="38" r="1.5" fill="#d4a843" opacity="0.6"/>
  <circle cx="25" cy="46" r="1.5" fill="#3ecf8e" opacity="0.6"/>
  <line x1="22" y1="42" x2="28" y2="38" stroke="#5585ff" stroke-width="0.7" opacity="0.35"/>
  <line x1="28" y1="38" x2="25" y2="46" stroke="#d4a843" stroke-width="0.7" opacity="0.3"/>
</svg>"""


# ═══════════════════════════════════════════════════════════════════════
# SESSION INIT
# ═══════════════════════════════════════════════════════════════════════
_DEFAULTS = {
    "authenticated": False, "page": "results",
    "queued_files": [], "extracted_papers": [],
    "synthesis_result": None, "processing_errors": [], "trigger_extract": False,
    # Credit system
    "access_granted": False, "access_error": "",
    "user_key": "", "user_credits": 0,
    "user_email": "", "user_row": None,
    "_credit_msg": None,
}
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v


# ═══════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION UTILS
# ═══════════════════════════════════════════════════════════════════════
def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        chunks = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        doc.close()
        return "\n\n".join(chunks)
    except Exception as e:
        raise RuntimeError(f"PDF parsing failed: {e}")

def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        raise RuntimeError(f"DOCX parsing failed: {e}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):   return _extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"): return _extract_text_from_docx(file_bytes)
    elif name.endswith(".txt"):  return file_bytes.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {filename}")

def smart_truncate(text: str, max_tokens: int = 12000) -> str:
    max_chars = max_tokens * 4
    if len(text) <= max_chars: return text
    patterns = [r"(abstract[\s\S]{0,3000})", r"(introduction[\s\S]{0,2000})",
                r"(method(?:ology|s)?[\s\S]{0,4000})", r"(result[\s\S]{0,3000})",
                r"(discussion[\s\S]{0,2000})", r"(conclusion[\s\S]{0,2000})"]
    found = []
    text_lower = text.lower()
    for pat in patterns:
        m = re.search(pat, text_lower)
        if m: found.append(text[m.start():m.end()])
    return ("\n\n".join(found) if found else text)[:max_chars]

def format_file_size(size_bytes: int) -> str:
    if size_bytes < 1024: return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024: return f"{size_bytes/1024:.1f} KB"
    return f"{size_bytes/(1024*1024):.1f} MB"


# ═══════════════════════════════════════════════════════════════════════
# OPENAI HELPERS
# ═══════════════════════════════════════════════════════════════════════
EXTRACTION_PROMPT = """You are an expert academic research analyst. Extract structured empirical information from the research paper text below.

CRITICAL RULES:
1. Only extract information explicitly present in the text. Do NOT invent or hallucinate.
2. If a field cannot be determined, use "Not specified" — never guess.
3. Be concise but complete. Findings must be directional.
4. For methodology, be specific.

Return a single JSON object with EXACTLY these keys:
- author_year, title, research_context, methodology, independent_variables, dependent_variable,
  control_variables, findings, theoretical_contributions, practical_contributions,
  strengths, limitations, citation_apa, citation_mla, citation_harvard

Return ONLY the JSON object. No preamble, no markdown fences.

Paper text:
{text}"""

SYNTHESIS_PROMPT = """You are a senior academic researcher producing a literature review section. You have extractions from {n_papers} papers.

Produce a DEEP, CITATION-RICH synthesis as JSON with EXACTLY these keys:
overall_summary, discussion_convergence, discussion_conflicts, discussion_methodology,
discussion_gaps, discussion_theory, discussion_implications,
common_findings (list), conflicting_results (list), dominant_methodology (string),
methodology_patterns (list), common_weaknesses (list), research_gaps (list),
underexplored_variables (list), future_directions (list)

Each discussion_* key: 3-6 sentences of flowing academic prose with inline citations like (Smith et al., 2021).
Lists: concise strings with citations.

CRITICAL: Base ALL content on the actual extractions. Return ONLY valid JSON.

Paper extractions:
{extractions}"""

def _get_openai_client():
    api_key = os.environ.get("OPENAI_API_KEY","")
    if not api_key:
        raise ValueError("OpenAI API key not set. Add it in the sidebar Settings field.")
    try:
        from openai import OpenAI
        return OpenAI(api_key=api_key)
    except ImportError:
        raise RuntimeError("`openai` package not installed.")

def extract_paper(text: str, filename: str = "") -> dict:
    client = _get_openai_client()
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role":"system","content":"Return valid JSON only."},
                  {"role":"user","content":EXTRACTION_PROMPT.format(text=text)}],
        temperature=0.1, max_tokens=1500, response_format={"type":"json_object"},
    )
    result = json.loads(response.choices[0].message.content)
    result["_source_file"] = filename; result["_status"] = "success"
    return result

def synthesize_papers(papers: list) -> dict:
    client = _get_openai_client()
    extractions_text = ""
    for i, p in enumerate(papers, 1):
        extractions_text += f"\n--- Paper {i} ---\n"
        for key in ["author_year","title","methodology","independent_variables","dependent_variable",
                    "control_variables","research_context","findings","theoretical_contributions","limitations","strengths"]:
            extractions_text += f"{key}: {p.get(key,'N/A')}\n"
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role":"system","content":"Return only valid JSON."},
                  {"role":"user","content":SYNTHESIS_PROMPT.format(n_papers=len(papers), extractions=extractions_text)}],
        temperature=0.25, max_tokens=4000, response_format={"type":"json_object"},
    )
    return json.loads(response.choices[0].message.content)


# ═══════════════════════════════════════════════════════════════════════
# EXPORTERS
# ═══════════════════════════════════════════════════════════════════════
DISPLAY_COLUMNS = [
    ("author_year","Author & Year"),("title","Title"),("research_context","Research Context"),
    ("methodology","Methodology"),("independent_variables","Independent Variables"),
    ("dependent_variable","Dependent Variable"),("control_variables","Control Variables"),
    ("findings","Key Findings"),("theoretical_contributions","Theoretical Contributions"),
    ("practical_contributions","Practical Contributions"),("strengths","Strengths"),("limitations","Limitations"),
]

def papers_to_csv(papers: list) -> bytes:
    output = io.StringIO()
    w = csv.writer(output)
    w.writerow([c[1] for c in DISPLAY_COLUMNS])
    for p in papers: w.writerow([p.get(c[0],"") for c in DISPLAY_COLUMNS])
    return output.getvalue().encode("utf-8")

def papers_to_excel(papers: list) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "EmpiricX Results"
    hf = Font(name="Calibri",bold=True,color="FFFFFF",size=10)
    hfill = PatternFill(start_color="07070d",end_color="07070d",fill_type="solid")
    ha = Alignment(horizontal="center",vertical="center",wrap_text=True)
    ca = Alignment(vertical="top",wrap_text=True)
    thin = Border(bottom=Side(style="thin",color="1b1b2f"),right=Side(style="thin",color="1b1b2f"))
    for ci,(key,label) in enumerate(DISPLAY_COLUMNS,1):
        cell = ws.cell(row=1,column=ci,value=label)
        cell.font=hf; cell.fill=hfill; cell.alignment=ha
    ws.row_dimensions[1].height = 36
    for ri,paper in enumerate(papers,2):
        for ci,(key,_) in enumerate(DISPLAY_COLUMNS,1):
            cell = ws.cell(row=ri,column=ci,value=paper.get(key,""))
            cell.alignment=ca; cell.border=thin; cell.font=Font(name="Calibri",size=9)
        fhex = "F5F5FA" if ri%2==0 else "FFFFFF"
        for ci in range(1,len(DISPLAY_COLUMNS)+1):
            ws.cell(row=ri,column=ci).fill = PatternFill(start_color=fhex,end_color=fhex,fill_type="solid")
    for ci,w in enumerate([18,30,22,18,22,18,18,35,30,30,22,25],1):
        ws.column_dimensions[get_column_letter(ci)].width = w
    ws.freeze_panes = "A2"
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

def synthesis_to_docx(synthesis: dict, papers: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc = Document()
    for s in doc.sections:
        s.top_margin=Inches(1.1); s.bottom_margin=Inches(1.1)
        s.left_margin=Inches(1.3); s.right_margin=Inches(1.3)
    doc.styles["Normal"].font.name="Georgia"; doc.styles["Normal"].font.size=Pt(11)
    C_DARK=RGBColor(0x07,0x07,0x0d); C_GOLD=RGBColor(0xd4,0xa8,0x43); C_GOLD_D=RGBColor(0x9a,0x7a,0x2a)
    C_BLUE=RGBColor(0x55,0x85,0xff); C_MID=RGBColor(0x44,0x44,0x66); C_MUTED=RGBColor(0x88,0x88,0xaa)
    def shd_para(para,hx):
        pPr=para._p.get_or_add_pPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); pPr.append(shd)
    def shd_cell(cell,hx):
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)
    def add_rule(color="d4a843",weight="4"):
        rp=doc.add_paragraph(); pPr=rp._p.get_or_add_pPr(); pb=OxmlElement("w:pBdr")
        bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),weight)
        bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),color); pb.append(bot); pPr.append(pb)
        rp.paragraph_format.space_after=Pt(0)
    def add_sh(num,title,sub=""):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(2)
        rn=p.add_run(f"{num}.  "); rn.font.name="Georgia"; rn.font.size=Pt(14); rn.font.bold=True; rn.font.color.rgb=C_GOLD
        rt=p.add_run(title); rt.font.name="Georgia"; rt.font.size=Pt(14); rt.font.bold=True; rt.font.color.rgb=C_DARK
        if sub:
            s=doc.add_paragraph(sub); s.paragraph_format.space_before=Pt(0); s.paragraph_format.space_after=Pt(8)
            s.runs[0].font.name="Arial"; s.runs[0].font.size=Pt(9); s.runs[0].font.italic=True; s.runs[0].font.color.rgb=C_MUTED
        add_rule("d4a84333","2")
    def add_prose(text):
        para=doc.add_paragraph(); para.paragraph_format.space_before=Pt(6); para.paragraph_format.space_after=Pt(8)
        para.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        run=para.add_run(text); run.font.name="Georgia"; run.font.size=Pt(11); run.font.color.rgb=C_MID
    def add_bullet(text):
        para=doc.add_paragraph(); para.paragraph_format.left_indent=Inches(0.3)
        para.paragraph_format.space_before=Pt(3); para.paragraph_format.space_after=Pt(3)
        rd=para.add_run("•  "); rd.font.name="Arial"; rd.font.size=Pt(10); rd.font.color.rgb=C_GOLD
        rt=para.add_run(str(text)); rt.font.name="Georgia"; rt.font.size=Pt(10.5); rt.font.color.rgb=C_MID
    # Cover
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    tr=tp.add_run("Empiri"); tr.font.name="Georgia"; tr.font.size=Pt(32); tr.font.bold=True; tr.font.color.rgb=C_DARK
    tx=tp.add_run("X"); tx.font.name="Georgia"; tx.font.size=Pt(32); tx.font.bold=True; tx.font.italic=True; tx.font.color.rgb=C_GOLD
    sub=doc.add_paragraph(); sub.add_run("Cross-Paper Literature Synthesis Report").font.color.rgb=C_MID
    meta=doc.add_paragraph()
    mr=meta.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}  ·  {len(papers)} paper(s) analysed")
    mr.font.name="Arial"; mr.font.size=Pt(9); mr.font.color.rgb=C_MUTED
    add_rule()
    summary=synthesis.get("overall_summary","")
    if summary:
        doc.add_paragraph(); lbl=doc.add_paragraph()
        lr=lbl.add_run("OVERVIEW"); lr.font.name="Arial"; lr.font.size=Pt(8); lr.font.bold=True; lr.font.color.rgb=C_GOLD
        sp=doc.add_paragraph(summary); sp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; sp.paragraph_format.space_after=Pt(12)
        shd_para(sp,"F0ECDE")
        if sp.runs: sp.runs[0].font.name="Georgia"; sp.runs[0].font.size=Pt(11.5); sp.runs[0].font.italic=True; sp.runs[0].font.color.rgb=RGBColor(0x33,0x22,0x00)
    doc.add_paragraph()
    for key,num,title,sub in [
        ("discussion_convergence","1","Convergence of Findings","Where the literature agrees."),
        ("discussion_conflicts","2","Conflicting Evidence","Divergences and contradictions."),
        ("discussion_methodology","3","Methodological Landscape","Critical appraisal of research designs."),
        ("discussion_theory","4","Theoretical Contributions","Frameworks invoked and theoretical advances."),
        ("discussion_gaps","5","Research Gaps","What remains unstudied or unresolved."),
        ("discussion_implications","6","Practical Implications","Actionable insights across all studies."),
    ]:
        prose=synthesis.get(key,"")
        if prose: add_sh(num,title,sub); add_prose(prose)
    doc.add_paragraph(); add_rule()
    qsh=doc.add_paragraph(); qr=qsh.add_run("Quick-Scan Summary")
    qr.font.name="Georgia"; qr.font.size=Pt(14); qr.font.bold=True; qr.font.color.rgb=C_DARK
    add_rule("5585ff","2")
    for key,title in [("common_findings","Common Findings"),("conflicting_results","Conflicting Results"),
                      ("methodology_patterns","Methodology Patterns"),("research_gaps","Research Gaps"),
                      ("common_weaknesses","Common Weaknesses"),("future_directions","Future Directions")]:
        items=synthesis.get(key,[])
        if items:
            sh=doc.add_paragraph(); sh.paragraph_format.space_before=Pt(10); sh.paragraph_format.space_after=Pt(4)
            sr2=sh.add_run(title); sr2.font.name="Arial"; sr2.font.size=Pt(10); sr2.font.bold=True; sr2.font.color.rgb=C_BLUE
            for item in items: add_bullet(item)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
    unexplored=synthesis.get("underexplored_variables",[])
    if unexplored:
        sh=doc.add_paragraph(); sh.paragraph_format.space_before=Pt(10)
        sr2=sh.add_run("Underexplored Variables"); sr2.font.name="Arial"; sr2.font.size=Pt(10); sr2.font.bold=True; sr2.font.color.rgb=C_BLUE
        vp=doc.add_paragraph(",  ".join(unexplored))
        if vp.runs: vp.runs[0].font.name="Georgia"; vp.runs[0].font.size=Pt(10.5); vp.runs[0].font.color.rgb=C_GOLD_D
    dom=synthesis.get("dominant_methodology","")
    if dom:
        sh=doc.add_paragraph(); sh.paragraph_format.space_before=Pt(10)
        sr2=sh.add_run("Dominant Methodology"); sr2.font.name="Arial"; sr2.font.size=Pt(10); sr2.font.bold=True; sr2.font.color.rgb=C_BLUE
        dp=doc.add_paragraph(dom); dp.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        if dp.runs: dp.runs[0].font.name="Georgia"; dp.runs[0].font.size=Pt(11); dp.runs[0].font.italic=True; dp.runs[0].font.color.rgb=C_MID
    if papers:
        doc.add_page_break()
        ph=doc.add_paragraph(); ph.paragraph_format.space_after=Pt(10)
        pr=ph.add_run("Analysed Papers"); pr.font.name="Georgia"; pr.font.size=Pt(16); pr.font.bold=True; pr.font.color.rgb=C_DARK
        add_rule()
        col_w=[Cm(3.5),Cm(5.5),Cm(3),Cm(2.8),Cm(4.2)]
        headers=["Author & Year","Title","Methodology","Context","Key Finding (brief)"]
        table=doc.add_table(rows=1,cols=5); table.style="Table Grid"
        for i,w in enumerate(col_w): table.columns[i].width=w
        hdr_row=table.rows[0]
        for i,h_text in enumerate(headers):
            cell=hdr_row.cells[i]; cell.width=col_w[i]
            run=cell.paragraphs[0].add_run(h_text)
            run.font.bold=True; run.font.name="Arial"; run.font.size=Pt(8.5); run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            shd_cell(cell,"07070d")
        for idx,paper in enumerate(papers):
            ft=paper.get("findings","") or ""
            short=ft[:100]+("…" if len(ft)>100 else "")
            vals=[paper.get("author_year","—"),paper.get("title","—"),paper.get("methodology","—"),paper.get("research_context","—")[:60],short]
            fill="F0F0FA" if idx%2==0 else "FFFFFF"
            row=table.add_row()
            for i,val in enumerate(vals):
                cell=row.cells[i]; cell.width=col_w[i]
                run=cell.paragraphs[0].add_run(val)
                run.font.name="Arial"; run.font.size=Pt(8); run.font.color.rgb=C_MID
                shd_cell(cell,fill)
    doc.add_paragraph()
    fp=doc.add_paragraph()
    fr=fp.add_run("Generated by EmpiricX — AI-Powered Empirical Research Intelligence  ·  v2.0")
    fr.font.name="Arial"; fr.font.size=Pt(8); fr.font.italic=True; fr.font.color.rgb=C_MUTED
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════
# API KEY BOOTSTRAP
# ═══════════════════════════════════════════════════════════════════════
try:
    key = st.secrets.get("OPENAI_API_KEY","")
    if key: os.environ["OPENAI_API_KEY"] = key
except Exception:
    pass


# ═══════════════════════════════════════════════════════════════════════
# LANDING PAGE + ACCESS GATE  (credit-key system)
# ═══════════════════════════════════════════════════════════════════════
if not st.session_state.access_granted:

    st.markdown("""
    <style>
    /* ── Hide sidebar on landing ── */
    [data-testid="stSidebar"],[data-testid="stSidebarCollapsedControl"]{display:none!important;}
    [data-testid="stAppViewContainer"]{padding:0!important;background:#07070d!important;}
    [data-testid="block-container"]{padding:0!important;max-width:100%!important;}
    section.main>div{padding:0!important;}
    footer,#MainMenu,[data-testid="stToolbar"]{display:none!important;}

    /* ── Ambient background ── */
    .lp-bg{position:fixed;inset:0;z-index:0;background:#07070d;}
    .lp-bg::before{content:'';position:absolute;inset:-60%;
      background:radial-gradient(ellipse 65% 50% at 50% -10%,rgba(212,168,67,.04) 0%,transparent 65%),
                 radial-gradient(ellipse 50% 65% at 100% 60%,rgba(155,126,255,.03) 0%,transparent 60%),
                 radial-gradient(ellipse 40% 40% at 0% 80%,rgba(212,168,67,.02) 0%,transparent 55%);
      animation:bgD 28s ease-in-out infinite alternate;}
    @keyframes bgD{0%{transform:scale(1) rotate(0)}100%{transform:scale(1.05) rotate(1.2deg)}}
    .lp-grid{position:fixed;inset:0;z-index:0;pointer-events:none;
      background-image:linear-gradient(rgba(212,168,67,.012) 1px,transparent 1px),
                       linear-gradient(90deg,rgba(212,168,67,.012) 1px,transparent 1px);
      background-size:72px 72px;
      mask-image:radial-gradient(ellipse 80% 80% at 50% 40%,black 20%,transparent 80%);}
    .lp-wrap{position:relative;z-index:10;}

    /* ── Animations ── */
    @keyframes fi{to{opacity:1;transform:translateY(0)}}
    .fi{opacity:0;transform:translateY(18px);animation:fi .6s cubic-bezier(.22,1,.36,1) forwards;}
    .d1{animation-delay:.06s}.d2{animation-delay:.14s}.d3{animation-delay:.22s}
    .d4{animation-delay:.30s}.d5{animation-delay:.38s}.d6{animation-delay:.46s}.d7{animation-delay:.54s}
    @keyframes pulse{0%,100%{opacity:1;transform:scale(1)}50%{opacity:.3;transform:scale(.85)}}
    .dot-live{display:inline-block;width:6px;height:6px;border-radius:50%;background:var(--gold);
      animation:pulse 2s ease-in-out infinite;vertical-align:middle;margin-right:7px;}

    /* ── Nav ── */
    .lp-nav{display:flex;align-items:center;justify-content:space-between;
      padding:18px clamp(20px,5vw,60px);border-bottom:1px solid var(--b1);
      background:rgba(7,7,13,.85);backdrop-filter:blur(24px);
      position:sticky;top:0;z-index:100;}
    .nav-brand{display:flex;align-items:center;gap:10px;font-family:var(--serif);
      font-weight:700;font-size:1.15rem;color:var(--t1);letter-spacing:-.02em;}
    .nav-logo{width:32px;height:32px;border-radius:8px;
      background:linear-gradient(135deg,rgba(212,168,67,.2),rgba(155,126,255,.2));
      border:1px solid rgba(212,168,67,.3);display:flex;align-items:center;
      justify-content:center;font-size:1rem;flex-shrink:0;}
    .brand-em{color:var(--gold);font-style:italic;}
    .nav-right{display:flex;align-items:center;gap:10px;}
    .nav-tag{font-family:var(--mono);font-size:.58rem;letter-spacing:.1em;text-transform:uppercase;
      padding:5px 13px;border-radius:100px;border:1px solid var(--b1);color:var(--t3);}
    .nav-tag-live{border-color:rgba(212,168,67,.28);color:var(--gold);background:rgba(212,168,67,.06);}
    @media(max-width:600px){.nav-right{display:none;}}

    /* ── Ticker ── */
    .lp-ticker{border-bottom:1px solid var(--b1);background:var(--bg-2);
      overflow:hidden;padding:9px 0;white-space:nowrap;position:relative;}
    .lp-ticker::before,.lp-ticker::after{content:'';position:absolute;top:0;bottom:0;width:80px;z-index:2;}
    .lp-ticker::before{left:0;background:linear-gradient(90deg,var(--bg-2),transparent);}
    .lp-ticker::after{right:0;background:linear-gradient(-90deg,var(--bg-2),transparent);}
    .ticker-track{display:inline-flex;animation:tick 40s linear infinite;}
    @keyframes tick{0%{transform:translateX(0)}100%{transform:translateX(-50%)}}
    .t-item{font-family:var(--mono);font-size:.58rem;color:var(--t3);letter-spacing:.12em;
      padding:0 28px;display:inline-flex;align-items:center;gap:7px;}
    .t-dot{color:var(--gold);font-size:.45rem;}

    /* ── Hero ── */
    .lp-hero{max-width:800px;margin:0 auto;
      padding:clamp(52px,8vw,96px) clamp(20px,5vw,48px) clamp(40px,6vw,72px);
      text-align:center;display:flex;flex-direction:column;align-items:center;}
    .hero-eyebrow{display:inline-flex;align-items:center;gap:8px;
      font-family:var(--mono);font-size:.6rem;letter-spacing:.2em;text-transform:uppercase;
      color:var(--gold);padding:6px 16px;border-radius:100px;
      border:1px solid rgba(212,168,67,.22);background:rgba(212,168,67,.05);margin-bottom:32px;}
    .hero-h1{font-family:var(--serif);font-weight:700;
      font-size:clamp(2.4rem,6vw,4.6rem);line-height:1.06;letter-spacing:-.025em;
      color:var(--t1);margin:0 0 24px;}
    .hero-h1 em{font-style:italic;color:var(--gold);}
    .hero-h1 .h1-muted{display:block;font-size:clamp(1.3rem,3vw,2.2rem);
      color:var(--t3);font-weight:400;margin-top:8px;}
    .hero-sub{font-family:var(--sans);font-size:clamp(.82rem,1.6vw,.96rem);
      color:var(--t2);line-height:1.9;max-width:560px;margin:0 0 48px;}
    .hero-sub strong{color:var(--t1);}
    .hero-pills{display:flex;flex-wrap:wrap;justify-content:center;gap:10px;margin-bottom:48px;}
    .h-pill{display:inline-flex;align-items:center;gap:8px;font-family:var(--mono);
      font-size:.62rem;letter-spacing:.06em;text-transform:uppercase;
      padding:8px 16px;border-radius:10px;border:1px solid var(--b2);
      background:var(--bg-3);color:var(--t2);transition:border-color .2s,color .2s;}
    .h-pill:hover{border-color:rgba(212,168,67,.3);color:var(--gold);}
    .hero-stats{display:flex;justify-content:center;flex-wrap:wrap;
      border:1px solid var(--b1);border-radius:14px;overflow:hidden;
      width:100%;max-width:560px;background:var(--bg-3);}
    .hs-block{flex:1;min-width:120px;padding:18px 16px;text-align:center;
      border-right:1px solid var(--b1);}
    .hs-block:last-child{border-right:none;}
    .hs-num{font-family:var(--serif);font-weight:700;font-size:1.4rem;color:var(--gold);line-height:1;}
    .hs-label{font-family:var(--mono);font-size:.54rem;color:var(--t3);
      text-transform:uppercase;letter-spacing:.14em;margin-top:6px;}

    /* ── Pricing section ── */
    .lp-pricing{max-width:900px;margin:0 auto;
      padding:clamp(32px,5vw,64px) clamp(20px,5vw,48px);}
    .section-head{text-align:center;margin-bottom:40px;}
    .section-label{font-family:var(--mono);font-size:.58rem;letter-spacing:.22em;
      text-transform:uppercase;color:var(--t3);margin-bottom:12px;display:block;}
    .section-title{font-family:var(--serif);font-weight:700;
      font-size:clamp(1.5rem,3vw,2rem);color:var(--t1);letter-spacing:-.02em;margin:0;}
    .section-title em{font-style:italic;color:var(--gold);}

    /* ── Price cards ── */
    .price-card{position:relative;background:var(--panel);border:1px solid var(--b1);
      border-radius:18px;padding:28px 24px 22px;
      transition:transform .25s,border-color .25s;overflow:hidden;}
    .price-card:hover{transform:translateY(-5px);border-color:var(--b2);}
    .price-card.featured{border-color:rgba(212,168,67,.3);
      background:linear-gradient(155deg,rgba(212,168,67,.03) 0%,rgba(155,126,255,.04) 100%);}
    .price-card::after{content:'';position:absolute;top:0;left:0;right:0;height:2px;
      opacity:0;transition:opacity .3s;
      background:linear-gradient(90deg,transparent,var(--gold) 50%,transparent);}
    .price-card.featured::after{opacity:1;}
    .price-card:hover::after{opacity:1;}
    .price-badge-pill{position:absolute;top:14px;right:14px;
      font-family:var(--mono);font-size:.5rem;letter-spacing:.12em;text-transform:uppercase;
      padding:3px 9px;border-radius:100px;}
    .badge-pop{background:rgba(212,168,67,.1);color:var(--gold);border:1px solid rgba(212,168,67,.28);}
    .badge-val{background:rgba(155,126,255,.1);color:var(--violet);border:1px solid rgba(155,126,255,.28);}
    .price-plan{font-family:var(--mono);font-size:.6rem;letter-spacing:.18em;text-transform:uppercase;
      color:var(--t3);margin-bottom:14px;}
    .price-amount{font-family:var(--serif);font-weight:700;font-size:2.5rem;
      color:var(--t1);line-height:1;letter-spacing:-.02em;}
    .price-curr{font-size:1rem;color:var(--t3);vertical-align:top;margin-top:8px;display:inline-block;}
    .price-credits{font-family:var(--mono);font-size:.7rem;color:var(--gold);margin:6px 0 18px;}
    .price-divider{height:1px;background:var(--b1);margin-bottom:16px;}
    .price-features{list-style:none;padding:0;margin:0 0 22px;}
    .price-features li{font-family:var(--mono);font-size:.64rem;color:var(--t2);
      padding:6px 0;display:flex;align-items:flex-start;gap:9px;
      border-bottom:1px solid rgba(255,255,255,.03);}
    .price-features li:last-child{border-bottom:none;}
    .pf-check{color:var(--gold);flex-shrink:0;margin-top:1px;}

    /* Streamlit link-button overrides for pricing cols */
    div[data-testid="column"]:nth-child(1) .stLinkButton a,
    div[data-testid="column"]:nth-child(3) .stLinkButton a{
      background:transparent!important;border:1px solid var(--b2)!important;
      color:var(--t1)!important;font-family:var(--serif)!important;
      font-weight:700!important;font-size:.72rem!important;border-radius:10px!important;
      padding:11px 0!important;width:100%!important;display:block!important;
      text-align:center!important;letter-spacing:.03em!important;
      transition:border-color .2s,color .2s!important;}
    div[data-testid="column"]:nth-child(1) .stLinkButton a:hover,
    div[data-testid="column"]:nth-child(3) .stLinkButton a:hover{
      border-color:rgba(212,168,67,.4)!important;color:var(--gold)!important;}
    div[data-testid="column"]:nth-child(2) .stLinkButton a{
      background:linear-gradient(135deg,#d4a843,#a87e2a)!important;border:none!important;
      color:#0a0700!important;font-family:var(--serif)!important;font-weight:700!important;
      font-size:.72rem!important;border-radius:10px!important;padding:11px 0!important;
      width:100%!important;display:block!important;text-align:center!important;
      box-shadow:0 4px 20px rgba(212,168,67,.32)!important;
      transition:box-shadow .2s,transform .2s!important;}
    div[data-testid="column"]:nth-child(2) .stLinkButton a:hover{
      box-shadow:0 7px 28px rgba(212,168,67,.52)!important;transform:translateY(-2px)!important;}
    .stLinkButton{margin:0!important;}

    /* ── Access gate ── */
    .lp-gate{max-width:540px;margin:0 auto;
      padding:clamp(16px,4vw,40px) clamp(20px,5vw,48px) clamp(48px,7vw,80px);}
    .gate-card{background:var(--panel);border:1px solid var(--b2);border-radius:20px;
      padding:clamp(28px,5vw,44px) clamp(24px,5vw,44px);
      position:relative;overflow:hidden;text-align:center;}
    .gate-card::before{content:'';position:absolute;top:0;left:0;right:0;height:1px;
      background:linear-gradient(90deg,transparent,var(--gold) 35%,var(--violet) 65%,transparent);
      opacity:.7;}
    .gate-lock{width:52px;height:52px;border-radius:14px;
      background:linear-gradient(135deg,rgba(212,168,67,.12),rgba(155,126,255,.12));
      border:1px solid rgba(212,168,67,.2);display:flex;align-items:center;
      justify-content:center;font-size:1.4rem;margin:0 auto 20px;}
    .gate-title{font-family:var(--serif);font-weight:700;font-size:1.5rem;
      color:var(--t1);letter-spacing:-.02em;margin-bottom:10px;}
    .gate-desc{font-family:var(--mono);font-size:.7rem;
      color:var(--t2);line-height:1.85;margin-bottom:28px;}
    .gate-desc strong{color:var(--t1);}
    .key-format{display:inline-flex;align-items:center;gap:8px;
      font-family:var(--mono);font-size:.6rem;color:var(--t3);
      background:rgba(255,255,255,.03);border:1px solid var(--b1);
      border-radius:8px;padding:7px 16px;margin-bottom:28px;}
    .kf-icon{color:var(--gold);}
    .gate-links{margin-top:20px;text-align:center;font-family:var(--mono);
      font-size:.6rem;color:var(--t3);line-height:2.4;
      display:flex;flex-wrap:wrap;justify-content:center;align-items:center;gap:4px 2px;}
    .gate-links a{color:var(--gold);text-decoration:none;}
    .gate-links a:hover{text-decoration:underline;text-underline-offset:2px;}
    .gate-sep{color:var(--b2);margin:0 6px;}
    .err-msg{margin-top:12px;padding:10px 14px;
      background:rgba(240,89,89,.06);border:1px solid rgba(240,89,89,.22);
      border-radius:10px;font-family:var(--mono);font-size:.65rem;color:var(--red);
      display:flex;align-items:center;gap:9px;text-align:left;}

    /* Input overrides */
    [data-testid="stTextInput"]>label{display:none!important;}
    [data-testid="stTextInput"]>div>div>input{
      background:rgba(255,255,255,.035)!important;border:1px solid rgba(255,255,255,.09)!important;
      border-radius:11px!important;color:var(--t1)!important;font-family:var(--mono)!important;
      font-size:.82rem!important;padding:14px 18px!important;letter-spacing:.1em!important;
      transition:border-color .2s,box-shadow .2s!important;}
    [data-testid="stTextInput"]>div>div>input:focus{
      border-color:rgba(212,168,67,.45)!important;
      box-shadow:0 0 0 3px rgba(212,168,67,.07)!important;outline:none!important;}
    [data-testid="stTextInput"]>div>div>input::placeholder{color:#2e3650!important;}
    [data-testid="baseButton-primary"]{
      background:linear-gradient(135deg,#d4a843 0%,#a87e2a 100%)!important;
      border:none!important;color:#04090f!important;
      font-family:var(--serif)!important;font-weight:700!important;
      font-size:.85rem!important;border-radius:11px!important;padding:14px!important;
      box-shadow:0 4px 22px rgba(212,168,67,.28)!important;
      transition:box-shadow .2s,transform .2s!important;}
    [data-testid="baseButton-primary"]:hover{
      box-shadow:0 8px 32px rgba(212,168,67,.48)!important;transform:translateY(-2px)!important;}

    /* ── Footer ── */
    .lp-footer{border-top:1px solid var(--b1);
      padding:20px clamp(20px,5vw,60px);display:flex;flex-wrap:wrap;
      justify-content:space-between;align-items:center;gap:10px;
      background:rgba(7,7,13,.92);}
    .footer-brand{font-family:var(--serif);font-size:.88rem;font-weight:700;
      color:var(--t3);letter-spacing:-.02em;}
    .footer-brand em{color:var(--gold);font-style:italic;}
    .footer-copy{font-family:var(--mono);font-size:.54rem;color:var(--t3);
      letter-spacing:.08em;text-transform:uppercase;opacity:.5;}
    </style>
    """, unsafe_allow_html=True)

    # ── Background layers ──
    st.markdown('<div class="lp-bg"></div><div class="lp-grid"></div>', unsafe_allow_html=True)

    # ── Nav ──
    st.markdown(f"""
    <nav class="lp-nav fi d1">
      <div class="nav-brand">
        <div class="nav-logo">◈</div>
        Empiric<em class="brand-em">X</em>
      </div>
      <div class="nav-right">
        <span class="nav-tag">Research Intelligence</span>
        <span class="nav-tag nav-tag-live"><span class="dot-live"></span>v2.0 Live</span>
      </div>
    </nav>
    """, unsafe_allow_html=True)

    # ── Ticker ──
    _ticker_items = [
        "Empirical Extraction", "Cross-Paper Synthesis", "Citation-Anchored Prose",
        "Research Gap Analysis", "Methodology Critique", "CSV & Excel Export",
        "Word Synthesis Report", "Convergence Mapping", "Conflict Detection",
        "Theoretical Landscape", "Future Directions", "AI-Powered Intelligence",
    ]
    _ticker_html = "".join(
        f'<span class="t-item"><span class="t-dot">◆</span>{item}</span>'
        for item in _ticker_items
    )
    st.markdown(f"""
    <div class="lp-ticker fi d1">
      <div class="ticker-track">{_ticker_html}{_ticker_html}</div>
    </div>
    """, unsafe_allow_html=True)

    # ── Hero ──
    st.markdown("""
    <div class="lp-hero">
      <div class="hero-eyebrow fi d2">
        <span class="dot-live"></span>
        AI-Powered Literature Intelligence
      </div>
      <h1 class="hero-h1 fi d2">
        Research <em>Synthesis</em>
        <span class="h1-muted">Re-imagined for rigorous scholars.</span>
      </h1>
      <p class="hero-sub fi d3">
        Upload any research papers — get <strong>structured empirical extractions</strong>,
        deep <strong>cross-paper synthesis</strong> with inline citations,
        and <strong>publication-ready Word reports</strong> in minutes, not hours.
      </p>
      <div class="hero-pills fi d4">
        <span class="h-pill"><span>◈</span> Empirical Extraction</span>
        <span class="h-pill"><span>◆</span> Cross-Paper Synthesis</span>
        <span class="h-pill"><span>◉</span> Citation-Anchored Prose</span>
        <span class="h-pill"><span>◇</span> Research Gap Analysis</span>
        <span class="h-pill"><span>▣</span> CSV / Excel / DOCX Export</span>
        <span class="h-pill"><span>✦</span> AI Model Powered</span>
      </div>
      <div class="hero-stats fi d5">
        <div class="hs-block">
          <div class="hs-num">PDF</div>
          <div class="hs-label">DOCX · TXT</div>
        </div>
        <div class="hs-block">
          <div class="hs-num">15+</div>
          <div class="hs-label">Fields Extracted</div>
        </div>
        <div class="hs-block">
          <div class="hs-num">AI</div>
          <div class="hs-label">Synthesis Engine</div>
        </div>
        <div class="hs-block">
          <div class="hs-num">DOCX</div>
          <div class="hs-label">Report Export</div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Pricing header ──
    st.markdown("""
    <div class="lp-pricing">
      <div class="section-head fi d5">
        <span class="section-label">Pay-As-You-Go Credits &middot; No Subscription</span>
        <p class="section-title">Simple, <em>transparent</em> pricing</p>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Pricing cards ──
    pc1, pc2, pc3 = st.columns(3, gap="small")
    with pc1:
        st.markdown("""
        <div class="price-card fi d3">
          <div class="price-plan">Starter</div>
          <div class="price-amount"><span class="price-curr">$</span>10</div>
          <div class="price-credits">10 Research Credits</div>
          <div class="price-divider"></div>
          <ul class="price-features">
            <li><span class="pf-check">&#x2713;</span>~7 papers extracted</li>
            <li><span class="pf-check">&#x2713;</span>1 synthesis run (up to 5 papers)</li>
            <li><span class="pf-check">&#x2713;</span>CSV &amp; Excel export</li>
            <li><span class="pf-check">&#x2713;</span>Word synthesis report</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Get Started →", "https://flutterwave.com/pay/YOUR_STARTER_LINK", width="stretch")

    with pc2:
        st.markdown("""
        <div class="price-card featured fi d4">
          <span class="price-badge-pill badge-pop">Most Popular</span>
          <div class="price-plan">Standard</div>
          <div class="price-amount"><span class="price-curr">$</span>30</div>
          <div class="price-credits">40 Research Credits</div>
          <div class="price-divider"></div>
          <ul class="price-features">
            <li><span class="pf-check">&#x2713;</span>~30 papers extracted</li>
            <li><span class="pf-check">&#x2713;</span>3–4 synthesis runs</li>
            <li><span class="pf-check">&#x2713;</span>CSV &amp; Excel export</li>
            <li><span class="pf-check">&#x2713;</span>Word synthesis report</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Buy Credits →", "https://flutterwave.com/pay/YOUR_STANDARD_LINK", width="stretch")

    with pc3:
        st.markdown("""
        <div class="price-card fi d5">
          <span class="price-badge-pill badge-val">Best Value</span>
          <div class="price-plan">Pro</div>
          <div class="price-amount"><span class="price-curr">$</span>80</div>
          <div class="price-credits">120 Research Credits</div>
          <div class="price-divider"></div>
          <ul class="price-features">
            <li><span class="pf-check">&#x2713;</span>~90 papers extracted</li>
            <li><span class="pf-check">&#x2713;</span>10+ synthesis runs</li>
            <li><span class="pf-check">&#x2713;</span>CSV &amp; Excel export</li>
            <li><span class="pf-check">&#x2713;</span>Word synthesis report</li>
          </ul>
        </div>
        """, unsafe_allow_html=True)
        st.link_button("Buy Credits →", "https://flutterwave.com/pay/YOUR_PRO_LINK", width="stretch")

    # ── Access gate ──
    st.markdown("""
    <div class="lp-gate fi d6">
      <div class="gate-card">
        <div class="gate-lock">&#x1F511;</div>
        <div class="gate-title">Enter Your Access Key</div>
        <div class="gate-desc">
          EmpiricX uses a <strong>credit-based, pay-as-you-go model</strong>.
          Each paper extraction costs <strong>1 credit</strong>. Synthesis costs
          <strong>1 credit per 5 papers</strong> (rounded up).
          Credits <strong>never expire</strong>.
        </div>
        <div class="key-format">
          <span class="kf-icon">◈</span>
          Key format: EMX-XXXX-XXXX-XXXX
        </div>
        <div style="text-align:left;font-family:var(--mono);font-size:.58rem;
                    text-transform:uppercase;letter-spacing:.16em;color:var(--t3);
                    margin-bottom:8px;">Access Key</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    _, gate_col, _ = st.columns([1, 2, 1])
    with gate_col:
        entered_key = st.text_input(
            "Access Key", type="password",
            placeholder="EMX-XXXX-XXXX-XXXX",
            label_visibility="collapsed",
        )
        unlock_btn = st.button("◈  Unlock EmpiricX", width="stretch", type="primary")

        if st.session_state.access_error:
            st.markdown(f"""
            <div class="err-msg"><span>&#x2715;</span> {st.session_state.access_error}</div>
            """, unsafe_allow_html=True)

        st.markdown("""
        <div class="gate-links">
          <a href="https://x.com/bayantx360" target="_blank" rel="noopener">&#x1F464; Get Access Key</a>
          <span class="gate-sep">|</span>
          <a href="mailto:bayantx360@gmail.com">&#x2699;&#xFE0F; Support</a>
        </div>
        """, unsafe_allow_html=True)

    if unlock_btn:
        if not entered_key:
            st.session_state.access_error = "Please enter your access key."
            st.rerun()
        else:
            with st.spinner("Verifying key…"):
                record = lookup_key(entered_key)
            if record is None:
                st.session_state.access_error = "Invalid access key. Please try again."
                st.rerun()
            elif record["credits"] <= 0:
                st.session_state.access_error = (
                    "Your credits have been exhausted. "
                    "Please purchase more credits to continue."
                )
                st.rerun()
            else:
                st.session_state.access_granted  = True
                st.session_state.access_error    = ""
                st.session_state.user_key        = record["key"]
                st.session_state.user_credits    = record["credits"]
                st.session_state.user_email      = record["email"]
                st.session_state.user_row        = record["row_index"]
                st.rerun()

    # ── Footer ──
    st.markdown("""
    <div class="lp-footer fi d7">
      <div class="footer-brand">Empiri<em>X</em></div>
      <div class="footer-copy">Research Intelligence Engine &middot; Credit-based access &middot; v2.0</div>
    </div>
    """, unsafe_allow_html=True)

    st.stop()


# ═══════════════════════════════════════════════════════════════════════
# EXTRACTION RUNNER
# ═══════════════════════════════════════════════════════════════════════
def run_extraction_from_queue():
    queued = st.session_state.get("queued_files",[])
    extracted = st.session_state.get("extracted_papers",[])
    extracted_names = {p.get("_source_file") for p in extracted}
    pending = [f for f in queued if f["name"] not in extracted_names]
    if not pending: return

    # ── Credit pre-flight check ──
    credits_needed = len(pending)
    if st.session_state.user_credits < credits_needed:
        st.error(
            f"⚠ Not enough credits. You need **{credits_needed}** credit(s) "
            f"({len(pending)} paper(s) × 1 credit each) but only have "
            f"**{st.session_state.user_credits}** remaining."
        )
        return

    total = len(pending); status = st.empty(); progress = st.progress(0.0)
    log_box = st.container(); errors = []
    status.info(f"🚀 Processing {total} paper(s) — {credits_needed} credit(s) will be deducted…")
    for i, finfo in enumerate(pending):
        fname = finfo["name"]; fobj = finfo.get("obj"); slot = log_box.empty(); t0 = time.time()
        slot.markdown(f'<div class="proc-card processing"><div class="spinner"></div><div><div class="proc-title">{fname}</div><div class="proc-sub">Parsing document...</div></div></div>', unsafe_allow_html=True)
        try:
            raw = fobj.read() if fobj else b""
            if not raw: raise ValueError("File is empty or could not be read.")
            text = extract_text(raw, fname)
            if len(text.strip()) < 150: raise ValueError("Too little text — file may be image-based.")
            text = smart_truncate(text, max_tokens=12000)
            slot.markdown(f'<div class="proc-card processing"><div class="spinner"></div><div><div class="proc-title">{fname}</div><div class="proc-sub">Extracting empirical insights...</div></div></div>', unsafe_allow_html=True)
            result = extract_paper(text, filename=fname)
            extracted.append(result)
            # Deduct 1 credit per successfully extracted paper
            new_credits = deduct_credits(st.session_state.user_row, st.session_state.user_credits, amount=1)
            st.session_state.user_credits = new_credits
            dur = round(time.time()-t0, 2)
            slot.markdown(f'<div class="proc-card success"><div>✔</div><div><div class="proc-title">{fname}</div><div class="proc-sub">Completed in {dur}s · {new_credits} credit(s) left</div></div></div>', unsafe_allow_html=True)
        except Exception as e:
            errors.append(f"**{fname}**: {e}")
            slot.markdown(f'<div class="proc-card error"><div>⚠</div><div><div class="proc-title">{fname}</div><div class="proc-sub">{e}</div></div></div>', unsafe_allow_html=True)
        progress.progress((i+1)/total)

    st.session_state["extracted_papers"] = extracted
    st.session_state["synthesis_result"] = None
    st.session_state["processing_errors"] = errors
    progress.empty()

    remaining = st.session_state.user_credits
    if remaining == 0:
        st.session_state._credit_msg = ("warn", "⚠ You just used your last credit.")
    elif remaining <= 3:
        st.session_state._credit_msg = ("warn", f"⚠ Only {remaining} credit(s) remaining.")

    if extracted: status.success(f"✅ {len(extracted)} paper(s) extracted and ready.")
    for e in errors: st.warning(e)


# ═══════════════════════════════════════════════════════════════════════
# PAGE: RESULTS
# ═══════════════════════════════════════════════════════════════════════
DISPLAY_COLS = [
    ("author_year","Author & Year"),("research_context","Research Context"),("methodology","Methodology"),
    ("independent_variables","IVs"),("dependent_variable","DV"),("findings","Key Findings"),
    ("strengths","Strengths"),("limitations","Limitations"),
]
FIELD_ICONS = {
    "research_context":"&#128205;","methodology":"&#9881;","independent_variables":"&#128204;",
    "dependent_variable":"&#127919;","control_variables":"&#128295;","findings":"&#128200;",
    "theoretical_contributions":"&#128161;","practical_contributions":"&#127981;",
    "strengths":"&#9989;","limitations":"&#9888;",
}

def _dblock(label, content):
    st.markdown(f'<div class="detail-block"><div class="detail-lbl">{label}</div><div class="detail-val">{content}</div></div>', unsafe_allow_html=True)

def page_results():
    papers = st.session_state.get("extracted_papers",[])
    st.markdown("""
    <div class="ph-wrap anim-up">
        <div class="ph-eye">Step 01 &middot; Analyse</div>
        <h1 class="ph-title">Extraction <span>Results</span></h1>
        <p class="ph-sub">Structured empirical data extracted from each paper. Inspect the full table or drill into any paper for detail.</p>
    </div>""", unsafe_allow_html=True)
    if not papers:
        st.markdown('<div class="empty-st"><span class="empty-st-icon">&#128230;</span><div class="empty-st-title">No papers extracted yet</div><div class="empty-st-desc">Upload papers using the sidebar, then click <strong>&#9889; Extract</strong> to begin.</div></div>', unsafe_allow_html=True)
        return
    methods = [p.get("methodology","") for p in papers if p.get("methodology") and p["methodology"]!="Not specified"]
    method_cnt = Counter(methods); top_m = method_cnt.most_common(1)[0][0][:24] if method_cnt else "—"
    dvs = [p.get("dependent_variable","") for p in papers if p.get("dependent_variable") and p["dependent_variable"]!="Not specified"]
    st.markdown(f"""
    <div class="m-row anim-up anim-up-d1">
        <div class="m-chip"><div class="m-val">{len(papers)}</div><div class="m-lbl">Papers Extracted</div></div>
        <div class="m-chip"><div class="m-val">{len(set(methods))}</div><div class="m-lbl">Unique Methods</div></div>
        <div class="m-chip"><div class="m-val">{len(set(dvs))}</div><div class="m-lbl">Outcome Variables</div></div>
        <div class="m-chip" style="min-width:180px"><div class="m-val" style="font-size:clamp(.9rem,2.2vw,1rem);padding-top:4px">{top_m}</div><div class="m-lbl">Top Methodology</div></div>
    </div>""", unsafe_allow_html=True)
    tab1, tab2 = st.tabs(["&#128202;  TABLE VIEW","&#128196;  PAPER DETAIL"])
    with tab1:
        rows = [{label:p.get(key,"—") for key,label in DISPLAY_COLS} for p in papers]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, height=min(620,90+65*len(papers)),
            column_config={"Author & Year":st.column_config.TextColumn(width="small"),"Research Context":st.column_config.TextColumn(width="medium"),
                           "Methodology":st.column_config.TextColumn(width="small"),"IVs":st.column_config.TextColumn(width="medium"),
                           "DV":st.column_config.TextColumn(width="small"),"Key Findings":st.column_config.TextColumn(width="large"),
                           "Strengths":st.column_config.TextColumn(width="medium"),"Limitations":st.column_config.TextColumn(width="medium")},
            hide_index=True)
    with tab2:
        options = [f"{i+1}. {p.get('author_year','Unknown')} — {p.get('_source_file','')}" for i,p in enumerate(papers)]
        idx = st.selectbox("Select paper", range(len(options)), format_func=lambda i: options[i], label_visibility="collapsed")
        p = papers[idx]
        st.markdown(f"""
        <div class="x-card x-card-gold anim-up" style="margin-bottom:1.25rem">
            <div style="font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--gold);margin-bottom:8px">{p.get("author_year","")}</div>
            <div style="font-family:var(--serif);font-size:1.2rem;color:var(--t1);margin-bottom:6px;font-style:italic;line-height:1.35">{p.get("title","Untitled")}</div>
            <div style="font-family:var(--mono);font-size:.68rem;color:var(--t3)">&#128196; {p.get("_source_file","")}</div>
        </div>""", unsafe_allow_html=True)
        labels = {"research_context":"Research Context","methodology":"Methodology","independent_variables":"Independent Variables",
                  "dependent_variable":"Dependent Variable","control_variables":"Control Variables","findings":"Key Findings",
                  "theoretical_contributions":"Theoretical Contributions","practical_contributions":"Practical Contributions",
                  "strengths":"Strengths","limitations":"Limitations"}
        c1,c2 = st.columns(2)
        with c1:
            for k in ["research_context","methodology","independent_variables","dependent_variable","control_variables"]:
                _dblock(f"{FIELD_ICONS.get(k,'')} {labels[k]}", p.get(k,"—"))
        with c2:
            for k in ["findings","theoretical_contributions","practical_contributions","strengths","limitations"]:
                _dblock(f"{FIELD_ICONS.get(k,'')} {labels[k]}", p.get(k,"—"))
        with st.expander("&#128218;  Citations — APA · MLA · Harvard"):
            for fmt,key in [("APA 7","citation_apa"),("MLA 9","citation_mla"),("Harvard","citation_harvard")]:
                st.markdown(f"**{fmt}**")
                st.code(p.get(key,"Not available"), language=None)
    st.markdown("---")
    c1,c2 = st.columns(2)
    with c1:
        if st.button("&#128279;  Run Cross-Paper Synthesis →", type="primary", use_container_width=True):
            st.session_state["page"]="synthesis"; st.rerun()
    with c2:
        if st.button("&#128229;  Export Data →", use_container_width=True):
            st.session_state["page"]="export"; st.rerun()


# ═══════════════════════════════════════════════════════════════════════
# PAGE: SYNTHESIS
# ═══════════════════════════════════════════════════════════════════════
def _syn_block(title, items, dot_cls, subtitle=""):
    if not items: return
    items_html = "".join(f'<div class="syn-item"><div class="syn-dot {dot_cls}"></div><div>{item}</div></div>' for item in items)
    sub_html = f'<div style="font-size:.74rem;color:var(--t3);margin-bottom:10px">{subtitle}</div>' if subtitle else ""
    st.markdown(f'<div class="syn-section"><div class="syn-head"><div class="syn-head-icon">{title.split()[0]}</div>{" ".join(title.split()[1:])}</div>{sub_html}{items_html}</div>', unsafe_allow_html=True)

def page_synthesis():
    papers = st.session_state.get("extracted_papers",[]); synthesis = st.session_state.get("synthesis_result")
    st.markdown("""
    <div class="ph-wrap anim-up">
        <div class="ph-eye">Step 02 &middot; Synthesize</div>
        <h1 class="ph-title">Cross-Paper <span>Synthesis</span></h1>
        <p class="ph-sub">A deep, citation-anchored literature synthesis — convergences, conflicts, methodological critique, theoretical landscape, and research gaps — produced in flowing scholarly prose.</p>
    </div>""", unsafe_allow_html=True)
    if not papers:
        st.markdown('<div class="empty-st"><span class="empty-st-icon">&#128279;</span><div class="empty-st-title">No papers loaded</div><div class="empty-st-desc">Upload and extract at least 2 papers to run synthesis.</div></div>', unsafe_allow_html=True)
        return
    if len(papers) < 2: st.warning("Add at least 2 papers for meaningful cross-paper synthesis."); return
    c1,c2 = st.columns([4,1])
    with c1:
        n_papers_syn = len(papers)
        syn_cost = synthesis_credit_cost(n_papers_syn)
        st.markdown(
            f'<p style="color:var(--t2);font-size:.88rem;margin:0">'
            f'<strong style="color:var(--gold)">{n_papers_syn}</strong> papers ready · '
            f'synthesis costs <strong style="color:var(--gold)">{syn_cost}</strong> credit(s) '
            f'({n_papers_syn} papers ÷ 5, rounded up)</p>',
            unsafe_allow_html=True
        )
    with c2:
        syn_disabled = st.session_state.user_credits < syn_cost
        run = st.button("&#128279; Run" if not synthesis else "&#128260; Re-run", type="primary", use_container_width=True, disabled=syn_disabled)
    if syn_disabled:
        st.caption(f"⚠ Need {syn_cost} credit(s) but only {st.session_state.user_credits} remaining.")
    if run:
        with st.spinner("Generating deep literature synthesis — this may take 30–60 seconds…"):
            try:
                result = synthesize_papers(papers)
                st.session_state["synthesis_result"] = result
                synthesis = result
                # Deduct credits for synthesis
                new_credits = deduct_credits(st.session_state.user_row, st.session_state.user_credits, amount=syn_cost)
                st.session_state.user_credits = new_credits
                if new_credits == 0:
                    st.session_state._credit_msg = ("warn", "⚠ You just used your last credit.")
                elif new_credits <= 3:
                    st.session_state._credit_msg = ("warn", f"⚠ Only {new_credits} credit(s) remaining.")
                st.rerun()
            except Exception as e: st.error(f"Synthesis failed: {e}"); return
    if not synthesis:
        st.markdown('<p style="color:var(--t3);font-size:.88rem;margin-top:1rem">Click <strong>Run</strong> above to generate the synthesis.</p>', unsafe_allow_html=True); return
    s = synthesis; st.markdown("---")
    summary = s.get("overall_summary","")
    if summary:
        st.markdown(f'<div class="syn-overview anim-up"><div class="syn-overview-label">Overview</div><div class="syn-overview-text">{summary}</div></div>', unsafe_allow_html=True)
    st.markdown('<div style="margin-bottom:1.25rem"><div style="font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:4px">&#128218; Literature Discussion</div><div style="font-size:.88rem;color:var(--t3)">Scholarly prose with inline citations — ready for your literature review.</div></div>', unsafe_allow_html=True)
    for key,num,title,subtitle in [
        ("discussion_convergence","1","Convergence of Findings","Where the literature agrees — patterns, consistencies, and shared insights."),
        ("discussion_conflicts","2","Conflicting Evidence","Divergences, contradictions, and contested findings."),
        ("discussion_methodology","3","Methodological Landscape","A critical appraisal of research designs and shared limitations."),
        ("discussion_theory","4","Theoretical Contributions","The frameworks invoked and advances this literature makes to theory."),
        ("discussion_gaps","5","Research Gaps","What remains unstudied, underexplored, or theoretically unresolved."),
        ("discussion_implications","6","Practical Implications","Actionable insights for managers and policymakers synthesised across all papers."),
    ]:
        prose = s.get(key,"")
        if prose:
            st.markdown(f"""
            <div class="syn-section lit-discussion anim-up">
                <div class="lit-section-header">
                    <div class="lit-section-num">{num}</div>
                    <div><div class="lit-section-title">{title}</div><div style="font-size:.78rem;color:var(--t3);margin-top:2px">{subtitle}</div></div>
                </div>
                <div class="lit-prose">{prose}</div>
            </div>""", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-family:var(--mono);font-size:.62rem;letter-spacing:2.5px;text-transform:uppercase;color:var(--t3);margin-bottom:1rem">&#9776; Quick-Scan Panels</div>', unsafe_allow_html=True)
    col1,col2 = st.columns(2)
    with col1:
        _syn_block("&#128200; Common Findings",    s.get("common_findings",[]),      "dot-gold",  "Recurring themes and results across papers")
        _syn_block("&#9881; Methodology Patterns",  s.get("methodology_patterns",[]), "dot-blue",  "Design choices and patterns observed")
        _syn_block("&#9888; Common Weaknesses",     s.get("common_weaknesses",[]),    "dot-red",   "Limitations shared across the literature")
    with col2:
        _syn_block("&#9889; Conflicting Results",   s.get("conflicting_results",[]),  "dot-red",   "Where studies disagree")
        _syn_block("&#128301; Research Gaps",       s.get("research_gaps",[]),        "dot-gold",  "Unstudied areas and open questions")
        _syn_block("&#128640; Future Directions",   s.get("future_directions",[]),    "dot-green", "Concrete recommendations for future research")
    unexplored = s.get("underexplored_variables",[])
    if unexplored:
        tags = "".join(f'<span class="var-tag">{v}</span>' for v in unexplored)
        st.markdown(f'<div class="syn-section" style="margin-top:.5rem"><div class="syn-head"><div class="syn-head-icon">&#128302;</div>Underexplored Variables</div><div style="margin-top:6px">{tags}</div></div>', unsafe_allow_html=True)
    dom = s.get("dominant_methodology","")
    if dom:
        st.markdown(f'<div class="dom-method-card"><div class="dom-method-icon">&#9881;</div><div><div class="dom-method-label">Dominant Methodology</div><div class="dom-method-val">{dom}</div></div></div>', unsafe_allow_html=True)
    st.markdown("---")
    if st.button("&#128229; Export Synthesis Report →", type="primary"):
        st.session_state["page"]="export"; st.rerun()


# ═══════════════════════════════════════════════════════════════════════
# PAGE: EXPORT
# ═══════════════════════════════════════════════════════════════════════
def page_export():
    papers = st.session_state.get("extracted_papers",[]); synthesis = st.session_state.get("synthesis_result")
    st.markdown("""
    <div class="ph-wrap anim-up">
        <div class="ph-eye">Step 03 &middot; Export</div>
        <h1 class="ph-title">Export <span>Results</span></h1>
        <p class="ph-sub">Download your extracted data and the full literature synthesis report — polished and ready for your paper, thesis, or briefing.</p>
    </div>""", unsafe_allow_html=True)
    if not papers:
        st.markdown('<div class="empty-st"><span class="empty-st-icon">&#128229;</span><div class="empty-st-title">Nothing to export yet</div><div class="empty-st-desc">Extract papers first to enable downloads.</div></div>', unsafe_allow_html=True)
        return
    st.markdown(f'<div class="detail-lbl" style="margin-bottom:1rem">&#128202; Paper Extraction Data &nbsp;&middot;&nbsp; {len(papers)} paper(s)</div>', unsafe_allow_html=True)
    c1,c2 = st.columns(2)
    with c1:
        st.markdown('<div class="exp-card"><span class="exp-icon">&#128203;</span><div class="exp-name">CSV Spreadsheet</div><div class="exp-desc">Universal format &middot; Opens in any tool &middot; All fields included</div></div>', unsafe_allow_html=True)
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        try: st.download_button("&#8595;  Download CSV", data=papers_to_csv(papers), file_name="empiricx_results.csv", mime="text/csv", use_container_width=True, key="dl_csv")
        except Exception as e: st.error(f"CSV error: {e}")
    with c2:
        st.markdown('<div class="exp-card"><span class="exp-icon">&#128209;</span><div class="exp-name">Excel Workbook</div><div class="exp-desc">Styled .xlsx &middot; Freeze panes &middot; Alternating rows</div></div>', unsafe_allow_html=True)
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        try: st.download_button("&#8595;  Download Excel", data=papers_to_excel(papers), file_name="empiricx_results.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True, key="dl_excel")
        except Exception as e: st.error(f"Excel error: {e}")
    st.markdown("---")
    st.markdown('<div class="detail-lbl" style="margin-bottom:1rem">&#128218; Synthesis Report &nbsp;&middot;&nbsp; Word Document</div>', unsafe_allow_html=True)
    if not synthesis:
        st.markdown('<div class="x-card" style="text-align:center;padding:2.5rem"><div style="font-size:2.5rem;margin-bottom:14px;opacity:.25">&#128196;</div><div style="font-size:.92rem;color:var(--t2);margin-bottom:18px;line-height:1.7">Run the Cross-Paper Synthesis first to generate the Word report.</div></div>', unsafe_allow_html=True)
        if st.button("&#128279; Go to Synthesis →"): st.session_state["page"]="synthesis"; st.rerun()
        return
    st.markdown('<div class="exp-card" style="text-align:left;max-width:520px"><div style="display:flex;align-items:center;gap:16px"><span style="font-size:2.4rem">&#128196;</span><div><div class="exp-name">Word Document (.docx)</div><div class="exp-desc" style="margin-top:6px">Full literature review report — scholarly overview, six deep discussion sections with inline citations, quick-scan summaries, and a reference table of all analysed papers.</div></div></div></div>', unsafe_allow_html=True)
    st.markdown('<div style="margin-top:12px"></div>', unsafe_allow_html=True)
    try: st.download_button("&#8595;  Download Synthesis Report (.docx)", data=synthesis_to_docx(synthesis,papers), file_name="empiricx_synthesis_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="dl_docx")
    except Exception as e: st.error(f"DOCX error: {e}")


# ═══════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown(f'<div class="sb-brand"><div class="sb-logomark">{LOGO_SVG}</div><span class="sb-name">Empiri<em>X</em></span></div><div class="sb-tag">Research Intelligence</div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    # ── Credit HUD ──────────────────────────────────────────────────────
    _credits_left = st.session_state.user_credits
    _cr_color = "#d4a843" if _credits_left > 5 else "#f5a623" if _credits_left > 1 else "#f05959"
    _cr_label  = "Credits remaining" if _credits_left > 1 else ("1 credit left!" if _credits_left == 1 else "No credits")
    _email_disp = st.session_state.user_email or "—"
    st.markdown(f"""
    <div style="background:var(--panel);border:1px solid var(--b1);border-radius:10px;
                padding:14px 16px;margin-bottom:16px;font-family:var(--mono);">
        <div style="font-size:.58rem;text-transform:uppercase;letter-spacing:.12em;
                    color:var(--t3);margin-bottom:8px;">Account</div>
        <div style="font-size:.68rem;color:var(--t3);margin-bottom:10px;
                    overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">{_email_disp}</div>
        <div style="display:flex;align-items:baseline;gap:6px;">
            <span style="font-family:var(--serif);font-size:1.6rem;font-weight:700;
                         color:{_cr_color};line-height:1;">{_credits_left}</span>
            <span style="font-size:.6rem;color:var(--t3);text-transform:uppercase;
                         letter-spacing:.1em;">{_cr_label}</span>
        </div>
        <div style="margin-top:8px;background:var(--b1);border-radius:2px;height:3px;">
            <div style="height:3px;border-radius:2px;background:{_cr_color};
                        width:{min(100, _credits_left * 5)}%;transition:width .4s;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Show deferred credit notification
    _cmsg = st.session_state.get("_credit_msg")
    if _cmsg:
        if _cmsg[0] == "warn":
            st.warning(_cmsg[1])
        st.session_state._credit_msg = None

    if _credits_left <= 0:
        st.error("⚠ No credits remaining. Please purchase more.")

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    # API key (if not set via secrets)
    if not os.environ.get("OPENAI_API_KEY"):
        st.markdown('<span class="sb-label">&#9881; OpenAI API Key</span>', unsafe_allow_html=True)
        api_key_input = st.text_input("API Key", type="password", placeholder="sk-...", label_visibility="collapsed", key="api_key_input")
        if api_key_input:
            os.environ["OPENAI_API_KEY"] = api_key_input
            st.success("Key saved for this session.")
        st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    st.markdown('<span class="sb-label">&#128196; Upload Papers</span>', unsafe_allow_html=True)
    uploaded = st.file_uploader("Upload", type=["pdf","docx","txt"], accept_multiple_files=True, label_visibility="collapsed", key="sidebar_uploader")
    if uploaded:
        queued = st.session_state.get("queued_files",[]); existing = {f["name"] for f in queued}
        for f in uploaded:
            if f.name not in existing:
                queued.append({"name":f.name,"size":f.size,"obj":f}); existing.add(f.name)
        st.session_state["queued_files"] = queued

    queued = st.session_state.get("queued_files",[]); extracted = st.session_state.get("extracted_papers",[])
    extracted_names = {p.get("_source_file") for p in extracted}

    if queued:
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        for finfo in queued:
            _badge_cls = "sb-badge-ok" if finfo["name"] in extracted_names else "sb-badge-new"
            _badge_txt = "Done" if finfo["name"] in extracted_names else "Queued"
            st.markdown(f'<div class="sb-paper"><div class="sb-paper-icon">&#128196;</div><div style="min-width:0;flex:1"><div class="sb-paper-name">{finfo["name"]}</div><div class="sb-paper-meta">{format_file_size(finfo["size"])}</div></div><span class="sb-badge {_badge_cls}">{_badge_txt}</span></div>', unsafe_allow_html=True)
        st.markdown('<div style="margin-top:10px"></div>', unsafe_allow_html=True)
        pending = [f for f in queued if f["name"] not in extracted_names]
        if pending:
            n_pending = len(pending)
            ext_disabled = st.session_state.user_credits < n_pending
            if st.button(f"⚡  Extract {n_pending} paper{'s' if n_pending>1 else ''}", use_container_width=True, disabled=ext_disabled):
                st.session_state["trigger_extract"]=True; st.session_state["page"]="results"; st.rerun()
            if ext_disabled:
                st.caption(f"⚠ Need {n_pending} credit(s), only {st.session_state.user_credits} left.")
            else:
                st.caption(f"⚡ Costs {n_pending} credit(s) · {st.session_state.user_credits} remaining")
        if st.button("🗑  Clear all", use_container_width=True, type="secondary"):
            st.session_state["queued_files"]=[]; st.session_state["extracted_papers"]=[]; st.session_state["synthesis_result"]=None; st.rerun()

    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    st.markdown('<span class="sb-label">Navigation</span>', unsafe_allow_html=True)
    nav_map = {"📊  Results":"results","🔗  Synthesis":"synthesis","📥  Export":"export"}
    selected = st.radio("nav", list(nav_map.keys()),
                        index=list(nav_map.values()).index(st.session_state.get("page","results") if st.session_state.get("page") in nav_map.values() else "results"),
                        label_visibility="collapsed", key="main_nav")
    st.session_state["page"] = nav_map[selected]
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

    n_papers = len(st.session_state.get("extracted_papers",[])); syn_done = "✓" if st.session_state.get("synthesis_result") else "—"
    st.markdown(f'<div class="sb-stats"><div class="sb-stat"><div class="sb-stat-n">{n_papers}</div><div class="sb-stat-l">Papers</div></div><div class="sb-stat"><div class="sb-stat-n">{syn_done}</div><div class="sb-stat-l">Synthesis</div></div></div>', unsafe_allow_html=True)
    st.markdown('<div style="margin-top:auto;padding-top:1rem"></div>', unsafe_allow_html=True)
    if st.button("⏻  Sign out", use_container_width=True, type="secondary"):
        for _k in ["access_granted","user_key","user_credits","user_email","user_row",
                   "access_error","_credit_msg","authenticated",
                   "queued_files","extracted_papers","synthesis_result"]:
            st.session_state[_k] = _DEFAULTS.get(_k, False if _k == "access_granted" else "" if _k in ("user_key","user_email","access_error") else 0 if _k in ("user_credits",) else None if _k in ("user_row","_credit_msg","synthesis_result") else [])
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════
# ROUTER
# ═══════════════════════════════════════════════════════════════════════
if st.session_state.get("trigger_extract"):
    st.session_state["trigger_extract"] = False
    run_extraction_from_queue()

page = st.session_state.get("page","results")
if page == "results":    page_results()
elif page == "synthesis": page_synthesis()
elif page == "export":    page_export()
